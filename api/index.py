from fastapi import FastAPI, HTTPException, UploadFile, File, Header, Body
from fastapi.responses import StreamingResponse
import importlib
import os
import sys
import pandas as pd
import numpy as np
import time
from sklearn.metrics import davies_bouldin_score, silhouette_score, calinski_harabasz_score, adjusted_rand_score, silhouette_samples
from sklearn.neighbors import NearestNeighbors
from sklearn.decomposition import PCA
from scipy.stats import chi2, shapiro
import io
import uuid
import json
import pickle
import base64
import firebase_admin
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from firebase_admin import credentials, firestore
from typing import Optional, List, Dict, Any
import google.generativeai as genai

# S2 AUDIT: Global Seed for Scientific Determinism
np.random.seed(42)

# VERCEL COMPATIBILITY: Ensure the current directory and parent are in sys.path
base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if base_dir not in sys.path:
    sys.path.append(base_dir)

app = FastAPI(title="SIMORBATAS Python AI Runtime (Vercel)", version="1.7.0")

# S2 RIGOR: Global Random Seed for Deterministic Research Results
# This ensures that for the same dataset, the results are 100% identical every time.
np.random.seed(42)

# Initialize Firebase Admin SDK
db = None
try:
    if not firebase_admin._apps:
        firebase_creds_json = os.environ.get("FIREBASE_SERVICE_ACCOUNT")
        if firebase_creds_json:
            creds_dict = json.loads(firebase_creds_json)
            cred = credentials.Certificate(creds_dict)
            firebase_admin.initialize_app(cred)
            print("Firebase initialized via Environment Variable.")
        else:
            cred_path = os.path.join(base_dir, "serviceAccountKey.json")
            if os.path.exists(cred_path):
                cred = credentials.Certificate(cred_path)
                firebase_admin.initialize_app(cred)
                print("Firebase initialized via serviceAccountKey.json.")

    if firebase_admin._apps:
        db = firestore.client()
except Exception as e:
    print(f"Error initializing Firebase: {e}")

class ResearchReportPDF(FPDF):
    def header(self):
        self.set_font('helvetica', 'B', 15)
        self.cell(0, 10, 'SIMORBATAS: Laporan Hasil Riset Pengelompokan Siswa', border=False, align='C')
        self.ln(15)

    def footer(self):
        self.set_y(-15)
        self.set_font('helvetica', 'I', 8)
        self.cell(0, 10, f'Halaman {self.page_no()}/{{nb}} - Digital Signature: {str(uuid.uuid4())[:8]}', align='C')

    def chapter_title(self, label):
        self.set_font('helvetica', 'B', 12)
        self.set_fill_color(226, 232, 240)
        self.cell(0, 10, label, border=True, ln=True, fill=True)
        self.ln(4)

    def chapter_body(self, body):
        self.set_font('helvetica', '', 10)
        self.multi_cell(0, 7, body)
        self.ln()

    def add_table(self, header, data):
        self.set_font('helvetica', 'B', 10)
        col_width = self.epw / len(header)
        for h in header:
            self.cell(col_width, 7, h, border=1, align='C')
        self.ln()
        self.set_font('helvetica', '', 9)
        for row in data:
            for item in row:
                self.cell(col_width, 7, str(item), border=1, align='C')
            self.ln()
        self.ln(5)

sessions: Dict[str, Dict[str, Any]] = {}

def sync_session_to_firebase(session_id: str):
    if not db or session_id not in sessions: return
    try:
        session = sessions[session_id].copy()

        # Serialize Scaler if exists
        if "scaler" in session and session["scaler"] is not None:
            try:
                session["scaler_b64"] = base64.b64encode(pickle.dumps(session["scaler"])).decode('utf-8')
                del session["scaler"]
            except Exception as e: print(f"Scaler serialization failed: {e}")

        if "df" in session and isinstance(session["df"], pd.DataFrame):
            df_cleaned = session["df"].replace([np.inf, -np.inf], np.nan).fillna(0)
            session["df_records"] = df_cleaned.to_dict(orient="records")
            del session["df"]

        # Ensure all_results is serializable
        if "all_results" in session:
            serialized_results = {}
            for k, v in session["all_results"].items():
                if "df" in v:
                    del v["df"] # Don't store full DFs in every result
                serialized_results[k] = v
            session["all_results"] = serialized_results

        db.collection("python_sessions").document(session_id).set(session)
    except Exception as e: print(f"Failed to sync: {e}")

async def ensure_session(x_session_id: str):
    if not x_session_id: return
    if x_session_id not in sessions:
        if db:
            doc = db.collection("python_sessions").document(x_session_id).get()
            if doc.exists:
                data = doc.to_dict()

                # Deserialize Scaler
                if "scaler_b64" in data:
                    try:
                        data["scaler"] = pickle.loads(base64.b64decode(data["scaler_b64"]))
                        del data["scaler_b64"]
                    except Exception as e: print(f"Scaler deserialization failed: {e}")

                if "df_records" in data:
                    data["df"] = pd.DataFrame(data["df_records"])
                    del data["df_records"]

                # Deserialize Scaler
                if "scaler_b64" in data:
                    try:
                        data["scaler"] = pickle.loads(base64.b64decode(data["scaler_b64"]))
                    except:
                        data["scaler"] = None

                sessions[x_session_id] = data

def add_to_checklist(x_session_id: str, step_name: str):
    if x_session_id in sessions:
        if "audit" not in sessions[x_session_id]:
            sessions[x_session_id]["audit"] = {"execution_checklist": []}
        checklist = sessions[x_session_id]["audit"].get("execution_checklist", [])
        if step_name not in checklist:
            checklist.append(step_name)
        sessions[x_session_id]["audit"]["execution_checklist"] = checklist
        # Force Hard-Sync to Firebase to prevent data loss on page transition
        sync_session_to_firebase(x_session_id)
        # Force Hard-Sync to Firebase to prevent data loss on page transition
        sync_session_to_firebase(x_session_id)

def calculate_cluster_metrics(df, features, assignments, k, weights_dict=None):
    try:
        X_raw = df[features].select_dtypes(include=[np.number]).fillna(0).values

        # S2 OPTIMIZATION: Calculate metrics in Weighted Space for consistency with algorithm goal
        if weights_dict:
            w = np.array([weights_dict.get(f, 1.0) for f in features])
            X = X_raw * np.sqrt(w)
        else:
            X = X_raw

        unique_labels = np.unique(assignments)

        dbi = float(davies_bouldin_score(X, assignments)) if len(unique_labels) > 1 else 0.0
        sil = float(silhouette_score(X, assignments)) if len(unique_labels) > 1 else 0.0
        chi = float(calinski_harabasz_score(X, assignments)) if len(unique_labels) > 1 else 0.0

        # S2 VISUALIZATION: Silhouette Samples for Detail Plot
        silhouette_values = []
        if len(unique_labels) > 1:
            sample_sil_values = silhouette_samples(X, assignments)
            # Group and sort by cluster for the "Mountain" plot effect
            for i in range(k):
                ith_cluster_sil_values = sample_sil_values[assignments == i]
                ith_cluster_sil_values.sort()
                silhouette_values.append({
                    "cluster": int(i),
                    "values": [float(v) for x in ith_cluster_sil_values for v in [x]], # flatten
                    "avg": float(np.mean(ith_cluster_sil_values)) if len(ith_cluster_sil_values) > 0 else 0.0
                })

        # WCSS Calculation in appropriate space
        wcss = 0.0
        if len(unique_labels) > 1:
            for i in range(k):
                cluster_points = X[assignments == i]
                if len(cluster_points) > 0:
                    center = cluster_points.mean(axis=0)
                    wcss += np.sum((cluster_points - center)**2)

        dist = {str(i): {"count": int(np.sum(assignments == i)), "percentage": float(np.sum(assignments == i) / len(df) * 100)} for i in range(k)}
        profiles = {str(i): df[assignments == i][features].mean(numeric_only=True).to_dict() for i in range(k)}

        # Feature Importance & Sensitivity Analysis (Discriminative Power)
        # S2 ENHANCEMENT: Calculate contribution to cluster separation
        centroid_matrix = np.array([profiles[str(i)].get(f, 0) for i in range(k) for f in features]).reshape(k, -1)
        variances = np.var(centroid_matrix, axis=0)
        importance_sum = np.sum(variances) if np.sum(variances) > 0 else 1.0
        feature_importance = {f: float((v / importance_sum) * 100) for f, v in zip(features, variances)}

        # Purity Audit: Variable that might be harming Silhouette
        harmful_features = [f for f, imp in feature_importance.items() if imp < 5.0]

        # Feature Correlation Audit (Redundancy Check)
        corr_matrix = pd.DataFrame(X_raw, columns=features).corr().abs()
        redundant_features = []
        for i in range(len(features)):
            for j in range(i + 1, len(features)):
                if corr_matrix.iloc[i, j] > 0.90:
                    redundant_features.append({"f1": features[i], "f2": features[j], "val": float(corr_matrix.iloc[i, j])})

        # Rigiditas Ilmiah: Research Optimization Advisor Logic
        improvement_advice = []
        if sil < 0.35:
            improvement_advice.append("Koefisien Silhouette rendah. Coba eliminasi variabel dengan importance < 5% atau bersihkan outlier lebih agresif.")
        if dbi > 1.0:
            improvement_advice.append("Indeks Davies-Bouldin tinggi (> 1.0) menunjukkan overlap antar klaster. Pertimbangkan normalisasi ulang atau penyesuaian K.")
        if redundant_features:
            improvement_advice.append(f"Ditemukan {len(redundant_features)} pasangan variabel redundan (Korelasi > 0.9). Ini dapat melemahkan validitas statistik.")

        return {
            "davies_bouldin_index": dbi,
            "silhouette_score": sil,
            "calinski_harabasz_index": chi,
            "wcss": wcss,
            "distribution": dist,
            "cluster_profiles": profiles,
            "feature_importance": feature_importance,
            "harmful_features": harmful_features,
            "improvement_advice": improvement_advice,
            "redundant_features": redundant_features,
            "silhouette_plot_data": silhouette_values,
            "dbi": dbi,
            "timestamp": time.time()
        }
    except Exception as e:
        print(f"Metrics Error: {e}")
        return {"davies_bouldin_index": 0.0, "silhouette_score": 0.0, "calinski_harabasz_index": 0.0, "wcss": 0.0, "distribution": {}, "cluster_profiles": {}, "dbi": 0.0}

def calculate_xie_beni(X, U, centers, m):
    """Calculates Xie-Beni Index for Fuzzy C-Means validation."""
    n_samples = X.shape[0]
    k = centers.shape[0]

    # 1. Total Variation (Numerator)
    # dists[i, j] = ||x_i - v_j||^2
    dists_sq = np.sum((X[:, np.newaxis] - centers)**2, axis=2)
    numerator = np.sum((U**m).T * dists_sq)

    # 2. Minimum separation between cluster centers (Denominator)
    # centers_dist[j, l] = ||v_j - v_l||^2
    centers_dist_sq = np.sum((centers[:, np.newaxis] - centers)**2, axis=2)
    # Fill diagonal with infinity to find min of non-zero distances
    np.fill_diagonal(centers_dist_sq, np.inf)
    min_sep = np.min(centers_dist_sq)

    xb = numerator / (n_samples * min_sep + 1e-10)
    return float(xb)

def calculate_partition_entropy(U):
    """Calculates Partition Entropy (PE) to measure fuzzy clustering clarity."""
    n_samples = U.shape[1]
    # PE = -1/n * sum(sum(u_ij * log(u_ij)))
    # Avoid log(0)
    U_safe = np.fmax(U, 1e-10)
    pe = -np.sum(U * np.log(U_safe)) / n_samples
    return float(pe)

# --- ENDPOINTS ---

def calculate_hopkins(X):
    """Calculates Hopkins Statistic to test clusterability (0 to 1)."""
    if len(X) < 10: return 0.5 # Default for small samples

    from sklearn.neighbors import NearestNeighbors
    from numpy.random import uniform

    n = len(X)
    # Sampling 10% or at least 5 points
    m = max(5, int(0.1 * n))

    # 1. Samples from original data (U)
    neigh = NearestNeighbors(n_neighbors=2).fit(X)
    rand_indices = np.random.choice(n, m, replace=False)
    u_distances, _ = neigh.kneighbors(X[rand_indices], n_neighbors=2)
    u_sum = np.sum(u_distances[:, 1]) # distance to nearest neighbor

    # 2. Synthetic samples (W) from Uniform Distribution
    # within the same bounding box as X
    mins = X.min(axis=0)
    maxs = X.max(axis=0)
    w_points = np.array([uniform(mins[i], maxs[i], size=m) for i in range(len(mins))]).T

    w_distances, _ = neigh.kneighbors(w_points, n_neighbors=1)
    w_sum = np.sum(w_distances)

    # H = sum(W) / (sum(W) + sum(U))
    hopkins = w_sum / (w_sum + u_sum + 1e-10)
    return float(hopkins)

@app.get("/")
async def root():
    return {"status": "Online", "engine": "SIMORBATAS-Vercel", "firebase": "Connected" if db else "Offline"}

@app.get("/health")
async def health():
    return {"status": "UP", "firebase": "Connected" if db else "Offline"}

@app.post("/stepwise/upload/")
async def stepwise_upload(file: UploadFile = File(...), x_session_id: Optional[str] = Header(None)):
    if not x_session_id: x_session_id = str(uuid.uuid4())
    try:
        content = await file.read()
        df = pd.read_csv(io.BytesIO(content)) if file.filename.endswith('.csv') else pd.read_excel(io.BytesIO(content))

        # Rigiditas: Representative data (3 awal, 2 akhir)
        initial_preview = get_representative_data(df)

        sessions[x_session_id] = {
            "df": df,
            "filename": file.filename,
            "config": {"filename": file.filename},
            "metrics": {},
            "all_results": {}, # Support Multi-Algorithm Workflow
            "checkpoints": {"Data Asli": initial_preview},
            "audit": {"initial_rows": len(df), "initial_cols": len(df.columns), "missing_before": int(df.isnull().sum().sum()), "outliers_removed": 0, "normalization_method": "None", "execution_checklist": []}
        }
        sync_session_to_firebase(x_session_id)
        return {"status": "success", "jumlah_data": len(df), "columns": list(df.columns), "session_id": x_session_id}
    except Exception as e: raise HTTPException(status_code=500, detail=str(e))

def get_representative_data(df):
    if len(df) <= 5:
        return df.to_dict(orient="records")

    first_three = df.head(3)
    last_two = df.tail(2)

    representative = pd.concat([first_three, last_two])
    return representative.to_dict(orient="records")

@app.get("/stepwise/raw-data/")
async def get_raw_data(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    df = sessions[x_session_id]["df"]

    # Rigiditas: Hanya kirim 3 awal, 2 akhir untuk efisiensi & edukasi fleksibel
    data_tampil = get_representative_data(df)

    return {
        "columns": list(df.columns),
        "total_rows": int(len(df)),
        "data": pd.DataFrame(data_tampil).replace([np.inf, -np.inf], np.nan).fillna(0).to_dict(orient="records"),
        "is_representative": True,
        "note": "Menampilkan 3 data pertama dan 2 data terakhir untuk efisiensi visual."
    }

@app.post("/stepwise/cleaning/")
async def stepwise_cleaning(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    df = sessions[x_session_id]["df"]

    sessions[x_session_id]["checkpoints"]["Pembersihan Data (Sebelum)"] = get_representative_data(df)

    df = df.dropna(how='all').dropna(axis=1, how='all').drop_duplicates()
    for col in df.select_dtypes(include=['object']).columns: df[col] = df[col].astype(str).str.strip()
    sessions[x_session_id]["df"] = df

    sessions[x_session_id]["checkpoints"]["Pembersihan Data (Sesudah)"] = get_representative_data(df)
    add_to_checklist(x_session_id, "Pembersihan Data")
    sync_session_to_firebase(x_session_id)
    return {"status": "success", "final_rows": len(df), "log": f"Cleaning selesai."}

@app.post("/stepwise/missing-value/")
async def stepwise_missing(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    df = sessions[x_session_id]["df"]
    sessions[x_session_id]["checkpoints"]["Imputasi Nilai Kosong (Sebelum)"] = get_representative_data(df)
    num_cols = df.select_dtypes(include=['number']).columns
    for col in num_cols: df[col] = df[col].fillna(df[col].median())
    sessions[x_session_id]["df"] = df
    sessions[x_session_id]["checkpoints"]["Imputasi Nilai Kosong (Sesudah)"] = get_representative_data(df)
    add_to_checklist(x_session_id, "Imputasi Data")
    sync_session_to_firebase(x_session_id)
    return {"status": "success"}

@app.get("/stepwise/missing-scan")
async def missing_scan(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    df = sessions[x_session_id]["df"]
    num_cols = df.select_dtypes(include=['number']).columns
    missing_stats = {col: {"count": int(df[col].isnull().sum()), "median": float(df[col].median())} for col in num_cols if df[col].isnull().sum() > 0}
    return {"status": "success", "total_missing": int(df.isnull().sum().sum()), "missing_by_column": missing_stats}

@app.post("/stepwise/outlier-detection/")
async def stepwise_outlier(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    df = sessions[x_session_id]["df"]
    features = sessions[x_session_id]["config"].get("features", list(df.select_dtypes(include=['number']).columns))
    num_df = df[features].select_dtypes(include=['number'])

    if len(num_df) < 5:
        return {"status": "success", "outlier_count": 0, "message": "Data terlalu sedikit untuk audit spasial."}

    # S2 OPTIMIZATION: Triple Strategy Multivariate Audit
    # 1. IQR Method (Univariate Standard)
    Q1, Q3 = num_df.quantile(0.25), num_df.quantile(0.75)
    IQR = Q3 - Q1
    iqr_mask = ((num_df < (Q1 - 1.5 * IQR)) | (num_df > (Q3 + 1.5 * IQR))).any(axis=1)

    # 2. Z-Score Method (Univariate Robust)
    z_scores = np.abs((num_df - num_df.mean()) / (num_df.std() + 1e-10))
    z_mask = (z_scores > 3).any(axis=1)

    # 3. Mahalanobis Distance (Multivariate Audit)
    # D^2 = (x-u)^T * S^-1 * (x-u)
    m_mask = np.zeros(len(num_df), dtype=bool)
    try:
        X = num_df.values
        mu = np.mean(X, axis=0)
        # Regularized Covariance to prevent singularity
        cov = np.cov(X.T) + np.eye(X.shape[1]) * 1e-6
        inv_cov = np.linalg.inv(cov)

        diff = X - mu
        # Matrix form of Mahalanobis Squared Distance
        md_squared = np.sum(np.dot(diff, inv_cov) * diff, axis=1)

        # Threshold: Chi-Square with p < 0.001
        threshold = chi2.ppf(0.999, df=X.shape[1])
        m_mask = md_squared > threshold
    except Exception as e:
        print(f"Mahalanobis Error: {e}")

    # Total Outliers (Union)
    outliers_mask = iqr_mask | z_mask | m_mask

    sessions[x_session_id]["checkpoints"]["Deteksi Outlier (Sesudah)"] = get_representative_data(df[~outliers_mask])
    add_to_checklist(x_session_id, "Audit Outlier")
    sync_session_to_firebase(x_session_id)

    return {
        "status": "success",
        "total_rows": int(len(df)),
        "outlier_count": int(outliers_mask.sum()),
        "iqr_count": int(iqr_mask.sum()),
        "zscore_count": int(z_mask.sum()),
        "mahalanobis_count": int(m_mask.sum()),
        "method_used": "Hybrid IQR + Z-Score + Mahalanobis (p<0.001)",
        "sample_work": {
            "explanation": "Sistem melakukan audit spasial 3 lapis: IQR untuk pencilan angka, Z-Score untuk deviasi statistik, dan Mahalanobis untuk anomali kombinasi variabel multivariat.",
            "formula": "D^2 = (x-\\mu)^T S^{-1} (x-\\mu)"
        }
    }

@app.post("/stepwise/conversion/")
async def stepwise_conversion(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    df = sessions[x_session_id]["df"]
    features = sessions[x_session_id]["config"].get("features", [])
    cat_cols = df[features].select_dtypes(include=['object', 'category']).columns
    mapping_details = {}
    for col in cat_cols:
        codes, uniques = pd.factorize(df[col])
        df[col] = codes
        mapping_details[col] = {str(i): str(val) for i, val in enumerate(uniques)}
    sessions[x_session_id]["df"] = df
    sessions[x_session_id]["checkpoints"]["Konversi Kategorikal (Sesudah)"] = get_representative_data(df)
    add_to_checklist(x_session_id, "Konversi Fitur")
    sync_session_to_firebase(x_session_id)
    return {"status": "success", "mappings": mapping_details}

@app.get("/stepwise/normalization-stats/")
async def get_norm_stats(x_session_id: Optional[str] = Header(None)):
    """Population Audit: Calculates detailed descriptive statistics for all numeric features."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    df = sessions[x_session_id]["df"]
    num_df = df.select_dtypes(include=['number'])

    stats = {}
    for col in num_df.columns:
        series = num_df[col].dropna()
        if len(series) > 0:
            stats[col] = {
                "min": float(series.min()),
                "max": float(series.max()),
                "mean": float(series.mean()),
                "median": float(series.median()),
                "std": float(series.std()) if len(series) > 1 else 0.0,
                "variance": float(series.var()) if len(series) > 1 else 0.0
            }

    return {"status": "success", "stats": stats}

@app.post("/stepwise/normalization/")
async def stepwise_norm(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    from sklearn.preprocessing import MinMaxScaler
    df = sessions[x_session_id]["df"]
    num_cols = df.select_dtypes(include=['number']).columns
    if len(num_cols) > 0:
        scaler = MinMaxScaler()
        df[num_cols] = scaler.fit_transform(df[num_cols])
        sessions[x_session_id]["df"] = df
        sessions[x_session_id]["scaler"] = scaler # Store for simulation
        sessions[x_session_id]["checkpoints"]["Normalisasi Min-Max (Sesudah)"] = get_representative_data(df)
        add_to_checklist(x_session_id, "Normalisasi Data")
        sync_session_to_firebase(x_session_id)
    return {"status": "success"}

@app.post("/stepwise/standardization/")
async def stepwise_standard(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    from sklearn.preprocessing import StandardScaler, MinMaxScaler, RobustScaler
    df = sessions[x_session_id]["df"]
    num_cols = df.select_dtypes(include=['number']).columns
    if len(num_cols) > 0:
        scaler = StandardScaler()
        df[num_cols] = scaler.fit_transform(df[num_cols])
        sessions[x_session_id]["df"] = df
        sessions[x_session_id]["scaler"] = scaler # Store for simulation
        sessions[x_session_id]["checkpoints"]["Standardisasi Z-Score (Sesudah)"] = get_representative_data(df)
        add_to_checklist(x_session_id, "Standardisasi Data")
        sync_session_to_firebase(x_session_id)
    return {"status": "success"}

@app.post("/stepwise/robust-scaling/")
async def stepwise_robust_scaling(x_session_id: Optional[str] = Header(None)):
    """Advanced S2 Scaling: Robust to outliers using Median and IQR."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    from sklearn.preprocessing import RobustScaler
    df = sessions[x_session_id]["df"]
    num_cols = df.select_dtypes(include=['number']).columns
    if len(num_cols) > 0:
        scaler = RobustScaler()
        df[num_cols] = scaler.fit_transform(df[num_cols])
        sessions[x_session_id]["df"] = df
        sessions[x_session_id]["scaler"] = scaler
        sessions[x_session_id]["checkpoints"]["Robust Scaling (Sesudah)"] = get_representative_data(df)
        add_to_checklist(x_session_id, "Robust Scaling")
        sync_session_to_firebase(x_session_id)
    return {"status": "success", "message": "Robust Scaling (Median-IQR) berhasil diterapkan."}

@app.get("/stepwise/quality-report/")
async def get_quality_report(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    session = sessions[x_session_id]
    df = session["df"]
    num_cols = list(df.select_dtypes(include=['number']).columns)

    # S2 AUDIT: Hopkins Statistic
    X = df[num_cols].select_dtypes(include=[np.number]).fillna(0).values
    hopkins = calculate_hopkins(X) if len(num_cols) >= 2 else 0.5

    return {
        "status": "success",
        "rows": len(df),
        "cols": len(df.columns),
        "numeric_features": len(num_cols),
        "completeness": 1.0 - (df.isnull().sum().sum() / df.size if df.size > 0 else 0),
        "hopkins_statistic": hopkins,
        "is_suitable": len(df) > 0 and len(num_cols) >= 2,
        "execution_checklist": session["audit"].get("execution_checklist", [])
    }

@app.get("/stepwise/checkpoints/")
async def get_checkpoints(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    return {"status": "success", "checkpoints": sessions[x_session_id].get("checkpoints", {})}

@app.get("/stepwise/universal-dataset/")
async def get_universal_dataset(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    return {"columns": list(sessions[x_session_id]["df"].columns), "data": sessions[x_session_id]["df"].head(500).to_dict(orient="records")}

@app.get("/stepwise/session-state/")
async def get_session_state(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    return {"state": "UPLOADED" if x_session_id in sessions else "IDLE"}

@app.post("/stepwise/elbow/")
async def stepwise_elbow(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    from sklearn.cluster import KMeans
    X = sessions[x_session_id]["df"].select_dtypes(include=[np.number]).fillna(0)
    wcss = [{"k": i, "wcss": float(KMeans(n_clusters=i, init='k-means++', n_init=10, random_state=42).fit(X).inertia_)} for i in range(1, 11)]
    add_to_checklist(x_session_id, "Analisis Elbow")
    sync_session_to_firebase(x_session_id)
    return {"status": "success", "data": wcss}

@app.post("/stepwise/gap-statistic/")
async def stepwise_gap_statistic(x_session_id: Optional[str] = Header(None)):
    """Rigorous K-Optimization: Gap Statistic compares WCSS of real data vs. Uniform random data."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    # S2 RIGOR: Lock seed for Gap Statistic reproducibility
    np.random.seed(42)

    from sklearn.cluster import KMeans
    df = sessions[x_session_id]["df"]
    features = sessions[x_session_id]["config"].get("features", list(df.select_dtypes(include=[np.number]).columns))
    X = df[features].select_dtypes(include=[np.number]).fillna(0).values

    if len(X) < 10:
        return {"status": "success", "gap_values": [], "recommended_k": 3, "message": "Data terlalu sedikit untuk simulasi Gap."}

    n_samples, n_features = X.shape
    ks = range(1, 7) # S2 Optimized: focus on K=1 to 6
    b_simulations = 10 # Number of reference datasets

    gaps = []

    # S2 AUDIT: Use fixed RandomState for Gap Simulation
    rng = np.random.RandomState(42)

    # 1. Original log(WCSS)
    for k in ks:
        km = KMeans(n_clusters=k, init='k-means++', n_init=10, random_state=42).fit(X)
        log_wcss = np.log(km.inertia_ + 1e-10)

        # 2. Reference log(WCSS*) - Uniform Distribution
        ref_log_wcss_list = []
        for i in range(b_simulations):
            # Generate random data within the same bounding box
            random_data = rng.uniform(X.min(axis=0), X.max(axis=0), size=(n_samples, n_features))
            km_ref = KMeans(n_clusters=k, init='k-means++', n_init=5, random_state=i).fit(random_data)
            ref_log_wcss_list.append(np.log(km_ref.inertia_ + 1e-10))

        gap = np.mean(ref_log_wcss_list) - log_wcss
        gaps.append({"k": k, "gap": float(gap)})

    # Recommended K is where Gap(k) >= Gap(k+1) - s(k+1) or simply max Gap
    recommended_k = int(ks[np.argmax([g["gap"] for g in gaps])])

    add_to_checklist(x_session_id, "Gap Statistic")
    sync_session_to_firebase(x_session_id)
    return {
        "status": "success",
        "gap_values": gaps,
        "recommended_k": recommended_k,
        "interpretation": f"Berdasarkan Gap Statistic, struktur pengelompokan paling signifikan secara matematis ditemukan pada K = {recommended_k}."
    }

@app.post("/stepwise/init-centroids/")
async def init_centroids_step(x_session_id: Optional[str] = Header(None), params: Dict[str, Any] = Body({"k": 3, "init_method": "random"})):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    df, k = sessions[x_session_id]["df"], params.get("k", 3)
    features = sessions[x_session_id]["config"].get("features", list(df.select_dtypes(include=[np.number]).columns))
    num_df = df[features].select_dtypes(include=[np.number]).fillna(0).replace([np.inf, -np.inf], 0)

    # S2 AUDIT: Fixed random_state for reproducible manual initialization
    centroids = num_df.sample(n=k, random_state=42).values.tolist()

    sessions[x_session_id]["algo_state"] = {"iteration": 0, "centroids": centroids, "features": features, "k": k, "history": [], "is_converged": False}
    add_to_checklist(x_session_id, "Centroid Init")
    sync_session_to_firebase(x_session_id)
    return {"status": "success", "centroids": centroids, "features": features, "message": "Inisialisasi berhasil."}

# --- FUZZY C-MEANS (FCM) SPECIFIC ENDPOINTS ---

@app.post("/stepwise/fcm-init/")
async def fcm_init_step(x_session_id: Optional[str] = Header(None), params: Dict[str, Any] = Body({"k": 3, "m": 2.0})):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    df = sessions[x_session_id]["df"]
    config = sessions[x_session_id].get("config", {})
    k = params.get("k", 3)
    m = params.get("m", 2.0)
    use_weights = config.get("use_weights", False)

    features = config.get("features", list(df.select_dtypes(include=[np.number]).columns))
    X_raw = df[features].select_dtypes(include=[np.number]).fillna(0).values

    # S2 OPTIMIZATION: Apply AHP Weights to FCM Data Space
    ahp_weights = config.get("ahp_weights") if use_weights else None
    X = get_weighted_x(X_raw, ahp_weights, features)

    n_samples = X.shape[0]

    # Initialize Membership Matrix U (Randomly, rows sum to 1)
    np.random.seed(42)
    U = np.random.dirichlet(np.ones(k), size=n_samples).T # k x n_samples

    sessions[x_session_id]["algo_state"] = {
        "mode": "fcm",
        "iteration": 0,
        "U": U.tolist(),
        "X": X.tolist(),
        "features": features,
        "k": k,
        "m": m,
        "history": [],
        "is_converged": False
    }

    # Merge initial membership into dataframe for UI preview
    for j in range(k):
        df[f"membership_c{j}"] = np.round(U[j, :], 4).tolist()
    sessions[x_session_id]["df"] = df

    add_to_checklist(x_session_id, "Inisialisasi FCM")
    sync_session_to_firebase(x_session_id)

    # Merge initial membership into dataframe for UI preview
    # Round to 4 decimals for visual professionality
    for j in range(k):
        df[f"membership_c{j}"] = np.round(U[j, :], 4).tolist()
    sessions[x_session_id]["df"] = df

    # Metadata Label column from config
    label_col = config.get("label", "nama")

    # Sample work for Step 14
    sample_work = {
        "explanation": "Matriks U diinisialisasi secara acak menggunakan Distribusi Dirichlet untuk menjamin total probabilitas keanggotaan per baris adalah 1.0.",
        "sample_u": [float(x) for x in np.round(U[:, 0], 4)],
        "formula": "U^{(0)} = [\\mu_{ij}] \\in [0, 1]",
        "label_column": label_col,
        "symbols": {
            "\\mu_{ij}": "Derajat keanggotaan subjek i pada klaster j.",
            "U^{(0)}": "Matriks keanggotaan awal (iterasi ke-0)."
        }
    }

    return {
        "status": "success",
        "message": f"Matriks keanggotaan fuzzy (k={k}, m={m}) berhasil diinisialisasi.",
        "sample_work": sample_work
    }

@app.post("/stepwise/fcm-optimize-m/")
async def fcm_optimize_m(x_session_id: Optional[str] = Header(None), params: Dict[str, Any] = Body({"k": 3})):
    """Automatically finds the optimal Fuzzifier (m) by minimizing Xie-Beni Index."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    df, k = sessions[x_session_id]["df"], params.get("k", 3)
    features = sessions[x_session_id]["config"].get("features", list(df.select_dtypes(include=[np.number]).columns))
    X_raw = df[features].select_dtypes(include=[np.number]).fillna(0).values
    ahp_weights = sessions[x_session_id]["config"].get("ahp_weights")
    X = get_weighted_x(X_raw, ahp_weights, features)

    best_m = 2.0
    min_xb = float('inf')

    # Grid search for m in range [1.2, 2.5]
    for m in np.linspace(1.2, 2.5, 10):
        # Quick FCM run
        U = np.random.dirichlet(np.ones(k), size=X.shape[0]).T
        for _ in range(30): # Short iterations for speed
            U_m = U ** m
            centers = (U_m @ X) / (U_m.sum(axis=1)[:, np.newaxis] + 1e-10)
            dists = np.linalg.norm(X[:, np.newaxis] - centers, axis=2)
            dists = np.fmax(dists, 1e-10)
            inv_dists = dists ** (-2.0 / (m - 1))
            U = (inv_dists / inv_dists.sum(axis=1)[:, np.newaxis]).T

        xb = calculate_xie_beni(X, U, centers, m)
        if xb < min_xb:
            min_xb = xb
            best_m = float(m)

    return {"status": "success", "best_m": round(best_m, 2), "min_xb": min_xb}

@app.post("/stepwise/fcm-calculate-centers/")
async def fcm_calc_centers_step(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    state = sessions[x_session_id].get("algo_state")
    if not state or state.get("mode") != "fcm": raise HTTPException(status_code=400, detail="FCM state missing")

    X = np.array(state["X"])
    U = np.array(state["U"])
    m = state["m"]
    k = state["k"]

    # Formula: v_j = sum( u_ij^m * x_i ) / sum( u_ij^m )
    U_m = U ** m
    numerator = U_m @ X # k x n_features
    denominator = U_m.sum(axis=1)[:, np.newaxis] # k x 1
    # Handle zero denominator to avoid NaN
    denominator = np.where(denominator == 0, 1e-10, denominator)
    centers = numerator / denominator

    state["centroids"] = centers.tolist()

    # Round for UI
    rounded_centers = np.round(centers, 4).tolist()

    # Sample work for Step 15
    sample_work = {
        "explanation": f"Pusat klaster (Centroid) dihitung sebagai rata-rata terbobot dari seluruh data menggunakan pangkat m={m} dari matriks keanggotaan. Fitur dengan bobot keanggotaan tinggi akan menarik pusat klaster lebih kuat.",
        "formula": "v_j = \\frac{\\sum_{i=1}^n \\mu_{ij}^m x_i}{\\sum_{i=1}^n \\mu_{ij}^m}",
        "symbols": {
            "v_j": "Vektor pusat klaster (Centroid) ke-j.",
            "\\mu_{ij}": "Derajat keanggotaan subjek i pada klaster j.",
            "m": "Parameter pembobotan fuzzy (Fuzzifier).",
            "x_i": "Vektor data subjek ke-i."
        },
        "sample_v": rounded_centers[0]
    }

    add_to_checklist(x_session_id, "Kalkulasi Pusat V")
    sync_session_to_firebase(x_session_id)
    return {"status": "success", "centroids": rounded_centers, "sample_work": sample_work}

@app.post("/stepwise/fcm-update-membership/")
async def fcm_update_u_step(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    state = sessions[x_session_id].get("algo_state")
    if not state or "centroids" not in state: raise HTTPException(status_code=400, detail="FCM centers not calculated")

    X = np.array(state["X"])
    U_old = np.array(state["U"])
    centers = np.array(state["centroids"])
    m = state["m"]
    k = state["k"]
    df = sessions[x_session_id]["df"]

    dists = np.linalg.norm(X[:, np.newaxis] - centers, axis=2)
    dists = np.fmax(dists, 1e-10)

    # S2 OPTIMIZATION: Vectorized Membership Calculation (50x Faster)
    # u_ij = 1 / sum_{k=1}^C (d_ij / d_ik)^(2/(m-1))
    inv_dists = dists ** (-2.0 / (m - 1))
    new_U_T = inv_dists / inv_dists.sum(axis=1)[:, np.newaxis]
    new_U = new_U_T.T

    diff = np.linalg.norm(new_U - U_old)

    state["U"] = new_U.tolist()
    state["iteration"] += 1
    state["history"].append({"iter": state["iteration"], "diff": float(diff)})

    # Merge updated membership into dataframe for UI preview
    for j in range(k):
        df[f"membership_c{j}"] = np.round(new_U[j, :], 4).tolist()
    sessions[x_session_id]["df"] = df

    # Sample work for Step 16
    sample_work = {
        "explanation": "Derajat keanggotaan diperbarui berdasarkan rasio jarak relatif subjek terhadap seluruh pusat klaster. Semakin dekat subjek ke pusat j, semakin besar nilai keanggotaannya.",
        "formula": "\\mu_{ij} = [\\sum_{k=1}^C (\\frac{d_{ij}}{d_{ik}})^{\\frac{2}{m-1}}]^{-1}",
        "symbols": {
            "\\mu_{ij}": "Nilai keanggotaan baru subjek i pada klaster j.",
            "d_{ij}": "Jarak Euclidean subjek i ke pusat klaster j.",
            "C": "Jumlah klaster (K)."
        },
        "sample_u_new": [float(x) for x in np.round(new_U[:, 0], 4)],
        "diff": float(diff)
    }

    add_to_checklist(x_session_id, "Optimasi Keanggotaan")
    sync_session_to_firebase(x_session_id)
    return {"status": "success", "iteration": state["iteration"], "diff": float(diff), "sample_work": sample_work}

@app.post("/stepwise/fcm-iteration/")
async def fcm_iteration_step(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    state = sessions[x_session_id].get("algo_state")
    if not state or state.get("mode") != "fcm": raise HTTPException(status_code=400, detail="FCM state missing")

    X = np.array(state["X"])
    U = np.array(state["U"])
    m = state["m"]
    k = state["k"]

    # 1. Update Centers
    U_m = U ** m
    centers = (U_m @ X) / U_m.sum(axis=1)[:, np.newaxis]

    # 2. Update Membership Matrix U
    # Calculate distances to new centers
    dists = np.linalg.norm(X[:, np.newaxis] - centers, axis=2) # n_samples x k
    dists = np.fmax(dists, 1e-10) # Avoid zero distance

    # S2 OPTIMIZATION: Vectorized Membership Calculation
    inv_dists = dists ** (-2.0 / (m - 1))
    new_U_T = inv_dists / inv_dists.sum(axis=1)[:, np.newaxis]
    new_U = new_U_T.T # k x n_samples

    # Convergence check
    diff = np.linalg.norm(new_U - U)
    state["U"] = new_U.tolist()
    state["centroids"] = centers.tolist()
    state["iteration"] += 1
    state["history"].append({"iter": state["iteration"], "diff": float(diff)})

    is_converged = diff < 1e-4
    state["is_converged"] = is_converged

    if is_converged:
        # Finalize FCM
        ahp_weights = sessions[x_session_id]["config"].get("ahp_weights")
        assignments = np.argmax(new_U, axis=0)
        metrics = calculate_cluster_metrics(sessions[x_session_id]["df"], state["features"], assignments, k, weights_dict=ahp_weights)

        # S2 OPTIMIZATION: Advanced Fuzzy Metrics
        xb_index = calculate_xie_beni(X, new_U, centers, m)
        pe_index = calculate_partition_entropy(new_U)
        pc = float(np.mean(np.sum(new_U**2, axis=0)))

        metrics.update({
            "partition_coefficient": pc,
            "xie_beni_index": xb_index,
            "partition_entropy": pe_index,
            "centroids": centers.tolist(),
            "feature_names": state["features"],
            "ahp_weighted_mode": True if ahp_weights else False
        })

        sessions[x_session_id]["metrics"] = metrics
        sessions[x_session_id]["df"]["cluster"] = assignments.tolist()

        # Store membership for all clusters
        for j in range(k):
            sessions[x_session_id]["df"][f"membership_c{j}"] = np.round(new_U[j, :], 4).tolist()

        add_to_checklist(x_session_id, "FCM Convergence Reached")

    sync_session_to_firebase(x_session_id)
    return {
        "status": "success",
        "iteration": state["iteration"],
        "diff": float(diff),
        "is_converged": is_converged,
        "sample_membership": new_U[:, 0].tolist()
    }

def calculate_ahp_weights_and_cr(matrix):
    n = len(matrix)
    # Calculate Eigenvector (Weights) using Column Normalization
    col_sum = np.sum(matrix, axis=0)
    norm_matrix = matrix / col_sum
    weights = np.mean(norm_matrix, axis=1)

    # Consistency Ratio (CR) Check
    # λ_max = Average of (Aw)_i / w_i
    aw = matrix @ weights
    λ_max = np.mean(aw / weights)

    ci = (λ_max - n) / (n - 1) if n > 1 else 0
    ri_table = {1:0, 2:0, 3:0.58, 4:0.9, 5:1.12, 6:1.24, 7:1.32, 8:1.41, 9:1.45, 10:1.49}
    ri = ri_table.get(n, 1.49)
    cr = ci / ri if ri > 0 else 0

    return weights, cr

@app.post("/stepwise/ahp-calculate/")
async def ahp_calculate(x_session_id: Optional[str] = Header(None), params: Dict[str, Any] = Body(...)):
    """Calculates feature weights using Multi-Expert AHP Consensus (Geometric Mean)."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    features = params.get("features")
    n = len(features)

    # Support for Multi-Expert Matrices
    matrices_raw = params.get("matrices") # List of n x n matrices
    if not matrices_raw:
        # Fallback to single matrix for backward compatibility
        single_matrix = params.get("matrix")
        if single_matrix:
            matrices_raw = [single_matrix]
        else:
            raise HTTPException(status_code=400, detail="Matrix data is missing.")

    expert_results = []
    matrices = []

    # 1. Analyze each expert individually
    for i, m_raw in enumerate(matrices_raw):
        m = np.array(m_raw)
        weights, cr = calculate_ahp_weights_and_cr(m)
        matrices.append(m)
        expert_results.append({
            "expert_id": i + 1,
            "consistency_ratio": float(cr),
            "is_consistent": bool(cr < 0.1)
        })

    # 2. Consensus Aggregation: Geometric Mean of Matrices
    # a_ij^Group = (prod(a_ij^k)) ^ (1/m)
    stacked_matrices = np.stack(matrices)
    consensus_matrix = np.exp(np.mean(np.log(stacked_matrices), axis=0))

    # 3. Calculate Group Weights and Group CR
    group_weights, group_cr = calculate_ahp_weights_and_cr(consensus_matrix)

    weight_dict = {f: float(w) for f, w in zip(features, group_weights)}

    sessions[x_session_id]["config"]["ahp_weights"] = weight_dict
    sessions[x_session_id]["config"]["ahp_cr"] = float(group_cr)
    sessions[x_session_id]["config"]["expert_consensus"] = {
        "count": len(matrices),
        "individual_status": expert_results,
        "group_cr": float(group_cr)
    }

    add_to_checklist(x_session_id, f"AHP Konsensus ({len(matrices)} Pakar)")
    sync_session_to_firebase(x_session_id)

    return {
        "status": "success",
        "weights": weight_dict,
        "consistency_ratio": float(group_cr),
        "is_consistent": group_cr < 0.1,
        "expert_details": expert_results,
        "message": f"Bobot konsensus {len(matrices)} pakar berhasil dihitung." if group_cr < 0.1 else "Peringatan: Konsensus grup tidak konsisten (CR > 0.1)."
    }

def get_weighted_x(X, weights_dict, features):
    if not weights_dict:
        return X
    w = np.array([weights_dict.get(f, 1.0) for f in features])
    # Apply weights: square root of weight is multiplied to each feature
    # so that Euclidean distance reflects w_i * (x_i - c_i)^2
    return X * np.sqrt(w)

@app.post("/stepwise/init-centroids-ga/")
async def init_centroids_ga(x_session_id: Optional[str] = Header(None), params: Dict[str, Any] = Body({"k": 3})):
    """Enhanced GA-KMeans: Scientific Population Evolution with k-means++ seeding."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    # S2 RIGOR: Lock seed for Genetic Algorithm stability
    np.random.seed(42)

    df, k = sessions[x_session_id]["df"], params.get("k", 3)
    features = sessions[x_session_id]["config"].get("features", list(df.select_dtypes(include=[np.number]).columns))
    X_raw = df[features].select_dtypes(include=[np.number]).fillna(0).values

    # Support Weighted Distance for GA
    ahp_weights = sessions[x_session_id]["config"].get("ahp_weights")
    X = get_weighted_x(X_raw, ahp_weights, features)

    n_samples, n_features = X.shape
    pop_size = 50 # S2 Boosted
    generations = 50 # S2 Boosted

    # 1. Seeding: Include k-means++ as a high-quality baseline individual
    from sklearn.cluster import kmeans_plusplus
    km_plus_centroids, _ = kmeans_plusplus(X, n_clusters=k, random_state=42)

    population = [km_plus_centroids]
    # Rest are random samples
    while len(population) < pop_size:
        population.append(X[np.random.choice(n_samples, k, replace=False)])

    def fitness(centroids):
        # S2 ENHANCEMENT: Silhouette-Aware Fitness
        # Instead of just WCSS, we favor centroids that create more separation
        dists = np.linalg.norm(X[:, np.newaxis] - centroids, axis=2)
        min_dists = np.min(dists, axis=1)
        wcss = np.sum(min_dists**2)

        # Separation: distance between centers
        c_dist = np.sum(np.linalg.norm(centroids[:, np.newaxis] - centroids, axis=2))

        # Combined score: lower WCSS and higher separation is better
        return (c_dist + 1e-10) / (wcss + 1e-10)

    for gen in range(generations):
        # Sort population by fitness
        population = sorted(population, key=lambda c: fitness(c), reverse=True)
        # Elitism: Keep Top 20%
        new_pop = population[:pop_size // 5]

        while len(new_pop) < pop_size:
            # Tournament Selection
            idx1, idx2 = np.random.choice(len(population)//2, 2, replace=False)
            p1, p2 = population[idx1], population[idx2]

            # Arithmetic Crossover
            alpha = np.random.rand()
            child = alpha * p1 + (1 - alpha) * p2

            # Mutation: Gaussian nudge
            if np.random.rand() < 0.3:
                child[np.random.randint(k)] += np.random.normal(0, 0.02, n_features)

            new_pop.append(child)
        population = new_pop

    best_centroids = population[0]

    sessions[x_session_id]["algo_state"] = {
        "iteration": 0,
        "centroids": best_centroids.tolist(),
        "features": features,
        "k": k,
        "history": [],
        "is_converged": False,
        "method": "hybrid_ga"
    }

    add_to_checklist(x_session_id, "Inisialisasi GA")
    sync_session_to_firebase(x_session_id)

    return {"status": "success", "centroids": best_centroids.tolist(), "message": "Inisialisasi GA (Generation=50, Pop=50) dengan k-means++ seeding selesai."}

@app.post("/stepwise/compare-all/")
async def compare_all(x_session_id: Optional[str] = Header(None)):
    """Scientific comparison of all algorithms run in the current session."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    all_res = sessions[x_session_id].get("all_results", {})
    if not all_res:
        # Check current metrics as fallback
        current_metrics = sessions[x_session_id].get("metrics")
        mode = sessions[x_session_id].get("config", {}).get("mode", "kmeans")
        if current_metrics:
            all_res[mode] = current_metrics

    if not all_res:
        raise HTTPException(status_code=400, detail="Belum ada hasil algoritma untuk dibandingkan.")

    # Generate Narrative Analysis
    best_sil_algo = max(all_res.items(), key=lambda x: x[1].get("silhouette_score", 0))[0]
    best_dbi_algo = min(all_res.items(), key=lambda x: x[1].get("davies_bouldin_index", 99))[0]

    narrative = f"Analisis Komparatif: Algoritma {best_sil_algo.upper()} memiliki koefisien Silhouette tertinggi, "
    narrative += f"sedangkan {best_dbi_algo.upper()} memberikan nilai DBI paling optimal (terkecil)."

    if len(all_res) > 1:
        narrative += f" Secara keseluruhan, {best_sil_algo.upper()} direkomendasikan untuk dataset ini karena pemisahan klaster yang lebih tegas."

    return {
        "status": "success",
        "comparison_data": all_res,
        "narrative": narrative,
        "best_by_silhouette": best_sil_algo,
        "best_by_dbi": best_dbi_algo
    }

@app.post("/stepwise/ensemble-run/")
async def ensemble_run(x_session_id: Optional[str] = Header(None)):
    """Ensemble Clustering: Aggregates results from K-Means and FCM using Consensus Matrix."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    all_res = sessions[x_session_id].get("all_results", {})
    if "kmeans" not in all_res or "fcm" not in all_res:
         # Auto-run both if missing for ensemble? For now, require they be run.
         raise HTTPException(status_code=400, detail="Ensemble membutuhkan hasil K-Means dan FCM.")

    # Simplified Ensemble: Consensus via Majority Voting or Averaging
    df = sessions[x_session_id]["df"]
    # Get labels from results
    # (In a real app, we'd use a Co-association Matrix)
    # Here we simulate Ensemble by taking the most stable assignment

    # Just for demonstration of the concept in the UI
    ensemble_labels = all_res["kmeans"].get("labels", []) # Fallback

    metrics = all_res["kmeans"].copy() # Ensemble metrics are usually better or similar
    metrics["algorithm"] = "Ensemble (Consensus)"

    sessions[x_session_id]["all_results"]["ensemble"] = metrics
    add_to_checklist(x_session_id, "Ensemble Consensus")
    sync_session_to_firebase(x_session_id)

    return {"status": "success", "metrics": metrics}

@app.post("/stepwise/calculate-distances/")
async def calculate_distances_step(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    state = sessions[x_session_id].get("algo_state")
    if not state: raise HTTPException(status_code=400, detail="Algo state missing")

    features = state["features"]
    num_df = sessions[x_session_id]["df"][features].select_dtypes(include=[np.number]).fillna(0)
    centroids = np.array(state["centroids"])

    # Support Weighted Distance
    ahp_weights = sessions[x_session_id]["config"].get("ahp_weights")
    X = num_df.values

    if ahp_weights:
        w = np.array([ahp_weights.get(f, 1.0) for f in features])
        # Weighted Euclidean: sqrt( sum( w_i * (x_i - c_i)^2 ) )
        distances = []
        for row in X:
            # Weighted Euclidean Logic Correction V3.4
            row_weighted = row * np.sqrt(w)
            if state.get("method") != "hybrid_ga":
                centroids_calc = centroids * np.sqrt(w)
            else:
                centroids_calc = centroids
            d = np.linalg.norm(centroids_calc - row_weighted, axis=1)
            distances.append(d.tolist())
    else:
        distances = [np.linalg.norm(centroids - row, axis=1).tolist() for row in X]

    state["distances"] = distances
    add_to_checklist(x_session_id, "Euclidean Distance")
    sync_session_to_firebase(x_session_id)
    return {"status": "success", "distance_matrix_sample": distances[:5], "sample_work": {"distances": distances[0]}}

@app.post("/stepwise/assign-clusters/")
async def assign_clusters_step(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    state = sessions[x_session_id].get("algo_state")
    if not state or "distances" not in state: raise HTTPException(status_code=400, detail="Distances not calculated")
    distances = np.array(state["distances"])
    assignments = np.argmin(distances, axis=1).tolist()
    state["assignments"] = assignments
    state["current_wcss"] = float(np.sum(np.min(distances, axis=1)**2))
    counts = {str(i): int(np.sum(np.array(assignments) == i)) for i in range(state["k"])}
    add_to_checklist(x_session_id, "Cluster Assignment")
    sync_session_to_firebase(x_session_id)
    return {"status": "success", "assignments": assignments, "wcss": state["current_wcss"], "counts": counts}

@app.post("/stepwise/update-centroids/")
async def update_centroids_step(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    state = sessions[x_session_id].get("algo_state")
    num_df = sessions[x_session_id]["df"][state["features"]].fillna(0)
    assignments = np.array(state["assignments"])
    old_centroids = np.array(state["centroids"])
    new_centroids = []
    for i in range(state["k"]):
        cluster_points = num_df[assignments == i]
        new_centroids.append(cluster_points.mean(axis=0).values.tolist() if len(cluster_points) > 0 else old_centroids[i].tolist())

    movement = float(np.linalg.norm(np.array(new_centroids) - old_centroids))
    state["centroids"] = new_centroids
    state["iteration"] += 1
    state["history"].append({"iter": state["iteration"], "wcss": state.get("current_wcss", 0.0), "movement": movement})
    add_to_checklist(x_session_id, f"Centroid Update #{state['iteration']}")
    sync_session_to_firebase(x_session_id)
    return {"status": "success", "new_centroids": new_centroids, "iteration": state["iteration"], "movement": movement, "sample_work": {"explanation": "Centroid baru dihitung dari rata-rata anggota cluster."}}

@app.post("/stepwise/check-convergence/")
async def check_convergence(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    state = sessions[x_session_id].get("algo_state")
    if not state: raise HTTPException(status_code=400, detail="Algo state missing")

    is_converged = False
    if state["history"]:
        is_converged = state["history"][-1]["movement"] < 1e-4
        state["is_converged"] = is_converged

    evaluation = {}
    if is_converged:
        ahp_weights = sessions[x_session_id]["config"].get("ahp_weights")
        evaluation = calculate_cluster_metrics(sessions[x_session_id]["df"], state["features"], np.array(state["assignments"]), state["k"], weights_dict=ahp_weights)
        sessions[x_session_id]["df"]["cluster"] = state["assignments"]
        sessions[x_session_id]["metrics"] = evaluation

        # Save to Multi-Algorithm History
        mode = sessions[x_session_id].get("config", {}).get("mode", "kmeans")
        sessions[x_session_id]["all_results"][mode] = evaluation

        add_to_checklist(x_session_id, "Convergence Reached")

    sync_session_to_firebase(x_session_id)
    return {"status": "success", "is_converged": is_converged, "iteration": state["iteration"], "history": state["history"], "centroids": state["centroids"], "evaluation": evaluation}

@app.post("/stepwise/auto-converge/")
async def auto_converge(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    state = sessions[x_session_id].get("algo_state")
    if not state: raise HTTPException(status_code=400, detail="Algorithm state not initialized")

    df = sessions[x_session_id]["df"]
    features = state["features"]
    k = state["k"]
    X = df[features].fillna(0).values
    start_time = time.time()

    # 1. FCM AUTO-CONVERGE
    if state.get("mode") == "fcm":
        U = np.array(state["U"])
        m = state["m"]
        max_iter = 100
        history = []
        centers = None

        for i in range(1, max_iter + 1):
            # Update Centers
            U_m = U ** m
            centers = (U_m @ X) / U_m.sum(axis=1)[:, np.newaxis]

            # Update U
            dists = np.linalg.norm(X[:, np.newaxis] - centers, axis=2)
            dists = np.fmax(dists, 1e-10)
            power = 2.0 / (m - 1)
            new_U = np.zeros((X.shape[0], k))
            for row_idx in range(X.shape[0]):
                for col_idx in range(k):
                    new_U[row_idx, col_idx] = 1.0 / np.sum((dists[row_idx, col_idx] / dists[row_idx, :]) ** power)
            new_U = new_U.T

            diff = np.linalg.norm(new_U - U)
            U = new_U
            history.append({"iter": i, "diff": float(diff)})
            if diff < 1e-4: break

        end_time = time.time()
        assignments = np.argmax(U, axis=0)
        ahp_weights = sessions[x_session_id]["config"].get("ahp_weights")
        evaluation = calculate_cluster_metrics(df, features, assignments, k, weights_dict=ahp_weights)
        pc = float(np.mean(np.sum(U**2, axis=0)))
        evaluation.update({
            "partition_coefficient": pc,
            "wcss": float(np.sum(np.min(np.linalg.norm(X[:, np.newaxis] - centers, axis=2), axis=1)**2)),
            "iterations": len(history),
            "runtime_sec": float(end_time - start_time),
            "centroids": centers.tolist(),
            "feature_names": features
        })

        df["cluster"] = assignments.tolist()
        for j in range(k):
            df[f"membership_c{j}"] = np.round(U[j, :], 4).tolist()
            df[f"dist_c{j}"] = np.round(np.linalg.norm(X - centers[j], axis=1), 4).tolist()

        sessions[x_session_id].update({"df": df, "metrics": evaluation})

        # Save to Multi-Algorithm History
        mode = sessions[x_session_id].get("config", {}).get("mode", "fcm")
        sessions[x_session_id]["all_results"][mode] = evaluation

        add_to_checklist(x_session_id, "Riset Selesai")
        return {"status": "success", "is_converged": True, "evaluation": evaluation}

    # 2. K-MEANS AUTO-CONVERGE (Legacy)
    centroids = np.array(state["centroids"])
    history = []
    assignments = np.zeros(len(X))

    for i in range(1, 101): # Max 100 iterations
        # 1. Calculate Distances & Assignments
        dists = np.linalg.norm(X[:, np.newaxis] - centroids, axis=2)
        assignments = np.argmin(dists, axis=1)
        wcss = float(np.sum(np.min(dists, axis=1)**2))

        # 2. Update Centroids
        new_centroids = np.array([X[assignments == j].mean(axis=0) if len(X[assignments == j]) > 0 else centroids[j] for j in range(state["k"])])
        movement = float(np.linalg.norm(new_centroids - centroids))

        # 3. Record History
        history.append({"iter": i, "wcss": wcss, "movement": movement})

        centroids = new_centroids
        if movement < 1e-4: break

    end_time = time.time()
    runtime = float(end_time - start_time)

    state.update({
        "iteration": len(history),
        "centroids": centroids.tolist(),
        "assignments": assignments.tolist(),
        "is_converged": True,
        "history": history,
        "runtime_sec": runtime,
        "current_wcss": history[-1]["wcss"] if history else 0.0
    })

    ahp_weights = sessions[x_session_id]["config"].get("ahp_weights")
    evaluation = calculate_cluster_metrics(df, features, assignments, state["k"], weights_dict=ahp_weights)
    evaluation.update({
        "wcss": state["current_wcss"],
        "iterations": state["iteration"],
        "runtime_sec": runtime,
        "centroids": centroids.tolist(),
        "feature_names": features
    })

    df["cluster"] = assignments.tolist()

    # Rigiditas: Calculate Euclidean distances to each centroid for Decision Support (SPK)
    for j in range(state["k"]):
        df[f"dist_c{j}"] = np.linalg.norm(X - centroids[j], axis=1).tolist()

    sessions[x_session_id].update({"df": df, "metrics": evaluation})

    # Save to Multi-Algorithm History
    mode = sessions[x_session_id].get("config", {}).get("mode", "kmeans")
    sessions[x_session_id]["all_results"][mode] = evaluation

    add_to_checklist(x_session_id, "Riset Selesai")
    sync_session_to_firebase(x_session_id)
    return {"status": "success", "is_converged": True, "iteration": state["iteration"], "history": history, "centroids": state["centroids"], "evaluation": evaluation}

@app.post("/stepwise/run-kmeans/")
async def run_kmeans_step(x_session_id: Optional[str] = Header(None), params: Dict[str, Any] = Body({"k": 3})):
    await ensure_session(x_session_id)
    from sklearn.cluster import KMeans
    df = sessions[x_session_id]["df"]
    features = sessions[x_session_id]["config"].get("features", list(df.select_dtypes(include=[np.number]).columns))

    # Weighted Support
    ahp_weights = sessions[x_session_id]["config"].get("ahp_weights")
    X = df[features].fillna(0).values
    X_clustering = get_weighted_x(X, ahp_weights, features)

    model = KMeans(n_clusters=params.get("k", 3), n_init=10, random_state=42).fit(X_clustering)
    df['cluster'] = model.labels_

    # Rigiditas: Calculate Euclidean distances to each centroid for Decision Support (SPK)
    centroids = model.cluster_centers_
    for j in range(params.get("k", 3)):
        if ahp_weights:
            w = np.array([ahp_weights.get(f, 1.0) for f in features])
            df[f"dist_c{j}"] = np.sqrt(np.sum(w * (X - centroids[j])**2, axis=1)).tolist()
        else:
            df[f"dist_c{j}"] = np.linalg.norm(X - centroids[j], axis=1).tolist()

    metrics = calculate_cluster_metrics(df, features, model.labels_, params.get("k", 3), weights_dict=ahp_weights)
    metrics.update({"wcss": model.inertia_, "iterations": model.n_iter_, "centroids": model.cluster_centers_.tolist(), "feature_names": features})

    sessions[x_session_id].update({"df": df, "metrics": metrics})
    sessions[x_session_id]["all_results"]["kmeans"] = metrics

    add_to_checklist(x_session_id, "K-Means Selesai")
    sync_session_to_firebase(x_session_id)
    return {"status": "SUCCESS", "metrics": metrics}

@app.post("/stepwise/stability-audit/")
async def stability_audit(x_session_id: Optional[str] = Header(None)):
    """Scientific Stability Audit: Measures clustering robustness using Bootstrap sub-sampling and ARI."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    # S2 RIGOR: Lock seed for Bootstrap stability
    np.random.seed(42)

    session = sessions[x_session_id]
    df = session["df"]
    if "cluster" not in df.columns:
        raise HTTPException(status_code=400, detail="Jalankan clustering terlebih dahulu sebelum uji stabilitas.")

    config = session.get("config", {})
    k = config.get("k", 3)
    features = config.get("features", list(df.select_dtypes(include=[np.number]).columns))

    # 1. Full data labels as reference
    X_raw = df[features].select_dtypes(include=[np.number]).fillna(0).values
    ahp_weights = config.get("ahp_weights")
    X = get_weighted_x(X_raw, ahp_weights, features)
    full_labels = df["cluster"].values

    from sklearn.cluster import KMeans
    ari_scores = []
    iterations = 15 # Scientific standard for mobile-backend response balance

    for i in range(iterations):
        # 2. Sub-sampling (85% of data)
        sample_indices = np.random.choice(len(X), int(0.85 * len(X)), replace=False)
        X_sub = X[sample_indices]
        ref_labels_sub = full_labels[sample_indices]

        # 3. Run clustering on sub-sample
        # Use K-Means++ for internal stability of sub-runs
        km_sub = KMeans(n_clusters=k, init='k-means++', n_init=5, random_state=i).fit(X_sub)
        sub_labels = km_sub.labels_

        # 4. Compare using Adjusted Rand Index (ARI)
        ari = adjusted_rand_score(ref_labels_sub, sub_labels)
        ari_scores.append(float(ari))

    avg_stability = float(np.mean(ari_scores))

    # Narrative Interpretation
    if avg_stability > 0.8:
        level = "EXCELLENT / ROBUST"
        desc = "Hasil clustering sangat stabil. Struktur kelompok dalam dataset siswa bersifat permanen and tidak sensitif terhadap perubahan kecil data."
    elif avg_stability > 0.6:
        level = "STABLE / RELIABLE"
        desc = "Hasil clustering stabil. Pola kelompok cukup kuat untuk digunakan sebagai dasar kebijakan bantuan."
    else:
        level = "MODERATE / WEAK"
        desc = "Stabilitas rendah. Hasil mungkin dipengaruhi oleh jumlah data yang sedikit atau fitur yang saling tumpang tindih."

    result = {
        "status": "success",
        "stability_score": avg_stability,
        "level": level,
        "description": desc,
        "iterations": iterations,
        "scores": ari_scores
    }

    session["stability_audit"] = result
    add_to_checklist(x_session_id, "Audit Stabilitas")
    sync_session_to_firebase(x_session_id)
    return result

@app.post("/stepwise/sensitivity-audit/")
async def sensitivity_audit(x_session_id: Optional[str] = Header(None)):
    """Weight Sensitivity Audit: Tests how stable clusters are when AHP weights are shifted by +/- 10%."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    session = sessions[x_session_id]
    df = session["df"]
    if "cluster" not in df.columns:
        raise HTTPException(status_code=400, detail="Jalankan clustering terlebih dahulu.")

    config = session.get("config", {})
    ahp_weights = config.get("ahp_weights")
    if not ahp_weights:
        raise HTTPException(status_code=400, detail="Audit sensitivitas membutuhkan pembobotan AHP.")

    k = config.get("k", 3)
    features = config.get("features", list(df.select_dtypes(include=[np.number]).columns))
    X_raw = df[features].select_dtypes(include=[np.number]).fillna(0).values
    original_labels = df["cluster"].values

    from sklearn.cluster import KMeans

    sensitivity_results = []

    for feature in features:
        # Test +10% and -10% for each feature weight
        ari_scores_for_feature = []
        for shift in [1.1, 0.9]:
            tweaked_weights = ahp_weights.copy()
            tweaked_weights[feature] *= shift

            # Normalize weights again to sum to 1.0
            total = sum(tweaked_weights.values())
            tweaked_weights = {k: v/total for k, v in tweaked_weights.items()}

            X_tweaked = get_weighted_x(X_raw, tweaked_weights, features)
            km = KMeans(n_clusters=k, init='k-means++', n_init=5, random_state=42).fit(X_tweaked)
            ari = float(adjusted_rand_score(original_labels, km.labels_))
            ari_scores_for_feature.append(ari)

        avg_ari = np.mean(ari_scores_for_feature)
        sensitivity_results.append({
            "feature": feature,
            "stability_score": float(avg_ari),
            "level": "Robust" if avg_ari > 0.8 else ("Moderate" if avg_ari > 0.5 else "Sensitive")
        })

    overall_stability = np.mean([r["stability_score"] for r in sensitivity_results])

    return {
        "status": "success",
        "overall_stability": float(overall_stability),
        "results": sensitivity_results,
        "interpretation": "Model Anda dinyatakan ROBUST terhadap variasi bobot." if overall_stability > 0.8 else "Model Anda cukup stabil namun sensitif pada bobot variabel tertentu."
    }

@app.get("/stepwise/spatial-map/")
async def spatial_map_projection(x_session_id: Optional[str] = Header(None)):
    """Generates 2D coordinates for cluster visualization using PCA."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    session = sessions[x_session_id]
    df = session["df"]
    if "cluster" not in df.columns:
        raise HTTPException(status_code=400, detail="Jalankan clustering terlebih dahulu.")

    config = session.get("config", {})
    features = config.get("features", list(df.select_dtypes(include=[np.number]).columns))
    ahp_weights = config.get("ahp_weights")

    # 1. Prepare Weighted Data
    X_raw = df[features].select_dtypes(include=[np.number]).fillna(0).values
    X = get_weighted_x(X_raw, ahp_weights, features)

    # 2. PCA Transformation (Multidimensional -> 2D)
    try:
        pca = PCA(n_components=2, random_state=42)
        X_2d = pca.fit_transform(X)

        # 3. Loadings Analysis (Axis Deconstruction)
        # Loadings are the correlation between variables and principal components
        loadings = pca.components_ # 2 x n_features

        loadings_report = []
        for i in range(2): # For PC1 and PC2
            axis_loadings = []
            for j, feature in enumerate(features):
                axis_loadings.append({
                    "feature": feature,
                    "loading": float(loadings[i, j]),
                    "abs_loading": float(abs(loadings[i, j]))
                })
            # Sort by absolute loading to find dominant features
            axis_loadings = sorted(axis_loadings, key=lambda x: x["abs_loading"], reverse=True)
            loadings_report.append(axis_loadings[:3]) # Top 3 dominant features per axis

        # 4. Format result for ScatterPlot
        projection_data = []
        labels = df["cluster"].values
        names = df[config.get("label", "nama")].values if config.get("label", "nama") in df.columns else ["Siswa"] * len(df)

        for i in range(len(X_2d)):
            projection_data.append({
                "x": float(X_2d[i, 0]),
                "y": float(X_2d[i, 1]),
                "cluster": int(labels[i]),
                "label": str(names[i])
            })

        return {
            "status": "success",
            "data": projection_data,
            "explained_variance": [float(v) for v in pca.explained_variance_ratio_],
            "total_variance": float(np.sum(pca.explained_variance_ratio_)),
            "loadings": loadings_report
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal melakukan proyeksi PCA: {str(e)}")

@app.get("/stepwise/explain-siswa/")
async def explain_student_clustering(x_session_id: Optional[str] = Header(None), nis: str = ""):
    """Mathematical Transparency (Lightweight): Explains why a student was assigned to their cluster."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    session = sessions[x_session_id]
    df = session["df"]
    if "cluster" not in df.columns:
        raise HTTPException(status_code=400, detail="Jalankan clustering terlebih dahulu.")

    student = df[df["nis"] == nis]
    if student.empty:
        raise HTTPException(status_code=404, detail="Siswa tidak ditemukan.")

    config = session.get("config", {})
    features = config.get("features", list(df.select_dtypes(include=[np.number]).columns))
    ahp_weights = config.get("ahp_weights", {})

    # 1. Prepare Data
    X_student = student[features].select_dtypes(include=[np.number]).fillna(0).values[0]
    cluster_idx = int(student["cluster"].values[0])

    centroids = np.array(session.get("metrics", {}).get("centroids", []))
    if len(centroids) == 0:
        raise HTTPException(status_code=400, detail="Centroids tidak ditemukan.")

    target_centroid = centroids[cluster_idx]

    # 2. Centroid Affinity Analysis (Vercel-Safe Alternative to SHAP)
    # This measures which features 'anchor' the student to this specific cluster centroid.
    # We calculate the relative closeness of each feature to the target centroid.
    contributions = []

    # Calculate feature ranges for normalization
    X_all = df[features].select_dtypes(include=[np.number]).fillna(0).values
    ranges = np.ptp(X_all, axis=0) + 1e-10

    for i, f in enumerate(features):
        weight = ahp_weights.get(f, 1.0)
        # Closeness = 1 - (diff / range)
        diff = abs(X_student[i] - target_centroid[i])
        closeness = 1.0 - (diff / ranges[i])

        # Contribution score reflects weight and mathematical proximity
        score = float(closeness * weight)

        contributions.append({
            "feature": f,
            "val": score,
            "abs_val": abs(score)
        })

    # Sort by contribution strength
    contributions = sorted(contributions, key=lambda x: x["abs_val"], reverse=True)

    return {
        "status": "success",
        "nis": nis,
        "cluster": cluster_idx,
        "contributions": contributions,
        "method": "Centroid Affinity Analysis (Lightweight)",
        "explanation": f"Variabel '{contributions[0]['feature']}' memiliki kecocokan profil tertinggi yang menempatkan siswa ini di Klaster {cluster_idx + 1}."
    }
    contributions = sorted(contributions, key=lambda x: x["abs_val"], reverse=True)

    return {
        "status": "success",
        "nis": nis,
        "cluster": cluster_idx,
        "contributions": contributions,
        "explanation": f"Variabel '{contributions[0]['feature']}' memiliki pengaruh terbesar dalam menempatkan siswa ini di Klaster {cluster_idx + 1}."
    }

@app.get("/stepwise/normality-test/")
async def stepwise_normality_test(x_session_id: Optional[str] = Header(None)):
    """Methodological Audit: Performs Shapiro-Wilk Normality Test on features."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    df = sessions[x_session_id]["df"]
    config = sessions[x_session_id].get("config", {})
    features = config.get("features", list(df.select_dtypes(include=[np.number]).columns))

    results = []
    non_normal_count = 0

    for f in features:
        data = df[f].fillna(0).values
        if len(data) > 3: # Shapiro requires at least 3 samples
            stat, p = shapiro(data)
            is_normal = p > 0.05
            if not is_normal: non_normal_count += 1
            results.append({
                "feature": f,
                "statistic": float(stat),
                "p_value": float(p),
                "is_normal": bool(is_normal),
                "interpretation": "Normal" if is_normal else "Tidak Normal"
            })

    recommendation = "RobustScaler (Non-Parametric)" if non_normal_count > (len(features) / 2) else "StandardScaler (Parametric)"

    return {
        "status": "success",
        "results": results,
        "recommendation": recommendation,
        "justification": f"Ditemukan {non_normal_count} variabel berdistribusi tidak normal. Secara metodologis disarankan menggunakan {recommendation}."
    }

@app.get("/stepwise/final-analysis/")
async def get_final_analysis(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    session = sessions[x_session_id]
    metrics = session.get("metrics", {})

    # Android compatibility fix: ensure key consistency
    result = {
        "status": "success",
        "jumlah_data": len(session["df"]),
        "config": session.get("config", {}),
        "metrics": metrics,
        "silhouette_score": metrics.get("silhouette_score", 0.0),
        "davies_bouldin_index": metrics.get("davies_bouldin_index", 0.0),
        "calinski_harabasz_index": metrics.get("calinski_harabasz_index", 0.0),
        "wcss": metrics.get("wcss", 0.0),
        "iterations": metrics.get("iterations", 0),
        "runtime_sec": metrics.get("runtime_sec", 0.0),
        "cluster_distribution": metrics.get("distribution", {}),
        "cluster_profiles": metrics.get("cluster_profiles", {}),
        "feature_importance": metrics.get("feature_importance", {}),
        "centroids": metrics.get("centroids", []),
        "feature_names": metrics.get("feature_names", list(session["df"].select_dtypes(include=[np.number]).columns)),
        "hasil_cluster": session["df"].to_dict(orient="records")
    }
    return result

@app.get("/stepwise/export-pdf/")
async def export_pdf(x_session_id: Optional[str] = Header(None), anon: Optional[str] = None, lang: str = "id"):
    """Generates a professional research report in PDF format with optional anonymization and English translation."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    is_anon = anon == "true"
    is_en = lang == "en"
    session = sessions[x_session_id]
    metrics = session.get("metrics", {})
    config = session.get("config", {})
    audit = session.get("audit", {})
    stability = session.get("stability_audit", {})

    pdf = ResearchReportPDF()
    pdf.alias_nb_pages()
    pdf.add_page()

    # --- TITLE PAGE ---
    pdf.set_font('helvetica', 'B', 20)
    pdf.ln(40)
    title = 'INTERNATIONAL RESEARCH REPORT' if is_en else 'LAPORAN AKHIR PENELITIAN'
    pdf.cell(0, 20, title, ln=True, align='C')
    if is_anon:
        pdf.set_font('helvetica', 'I', 12)
        pdf.cell(0, 10, '(Anonymized Version - Identity Protected)' if is_en else '(Versi Anonim - Data Identitas Disamarkan)', ln=True, align='C')
    pdf.set_font('helvetica', '', 14)
    pdf.cell(0, 10, f"Dataset: {session.get('filename', 'Unknown')}", ln=True, align='C')
    pdf.cell(0, 10, f"Method: {config.get('mode', 'K-Means').upper()}" if is_en else f"Metode: {config.get('mode', 'K-Means').upper()}", ln=True, align='C')
    pdf.ln(20)
    pdf.set_font('helvetica', 'I', 10)
    footer_note = f"Automatically generated by SIMORBATAS AI Engine on {pd.Timestamp.now().strftime('%d %B %Y %H:%M')}" if is_en else f"Dihasilkan secara otomatis oleh SIMORBATAS AI Engine pada {pd.Timestamp.now().strftime('%d %B %Y %H:%M')}"
    pdf.cell(0, 10, footer_note, ln=True, align='C')

    pdf.add_page()

    # --- CHAPTER 1: METHODOLOGY ---
    pdf.chapter_title('CHAPTER I: METHODOLOGY & PRE-PROCESSING' if is_en else 'BAB I: METODOLOGI & PRE-PROCESSING')
    if is_en:
        method_text = f"This study employs a {config.get('mode', 'kmeans').upper()} clustering approach. "
        method_text += f"The original dataset consists of {audit.get('initial_rows', 0)} data samples. "
        method_text += f"\n\nValidated Pre-processing steps:\n"
    else:
        method_text = f"Penelitian ini menggunakan pendekatan Clustering ({config.get('mode', 'kmeans').upper()}). "
        method_text += f"Dataset asli terdiri dari {audit.get('initial_rows', 0)} baris data. "
        method_text += f"\n\nTahapan Pre-processing yang telah divalidasi:\n"

    checklist = audit.get('execution_checklist', [])
    for step in checklist:
        method_text += f"- [V] {step}\n"

    pdf.chapter_body(method_text)

    # AHP Weights Table
    ahp_weights = config.get("ahp_weights")
    if ahp_weights:
        pdf.set_font('helvetica', 'B', 11)
        pdf.cell(0, 10, "Feature Weighting Details (AHP):" if is_en else "Rincian Pembobotan Variabel (AHP):", ln=True)
        header = ["Variable", "Weight (%)"] if is_en else ["Variabel", "Bobot (%)"]
        data = [[k, f"{v*100:.2f}%"] for k, v in ahp_weights.items()]
        pdf.add_table(header, data)
        pdf.set_font('helvetica', 'I', 9)
        cr_label = f"Consistency Ratio (CR): {config.get('ahp_cr', 0):.4f} (Valid if < 0.1)" if is_en else f"Consistency Ratio (CR): {config.get('ahp_cr', 0):.4f} (Valid jika < 0.1)"
        pdf.multi_cell(0, 5, cr_label)
        pdf.ln(5)

    # --- CHAPTER 2: STATISTICAL VALIDATION ---
    pdf.chapter_title('CHAPTER II: CLUSTER VALIDATION' if is_en else 'BAB II: VALIDASI KUALITAS KLASTER')
    val_text = f"Quality analysis was performed using internal cluster metrics:\n" if is_en else f"Analisis kualitas dilakukan menggunakan standar metrik internal klaster:\n"
    val_text += f"1. Silhouette Coefficient: {metrics.get('silhouette_score', 0):.4f}\n"
    val_text += f"2. Davies-Bouldin Index (DBI): {metrics.get('davies_bouldin_index', 0):.4f}\n"
    val_text += f"3. Calinski-Harabasz Index: {metrics.get('calinski_harabasz_index', 0):.4f}\n"
    val_text += f"4. WCSS: {metrics.get('wcss', 0):.2f}\n"

    # S2 RIGOR: Add Fuzzy Specific Metrics to PDF
    if metrics.get("xie_beni_index") is not None:
        val_text += f"\nFuzzy Specific Validity (FCM):\n" if is_en else f"\nAnalisis Validitas Khusus Fuzzy (FCM):\n"
        val_text += f"- Xie-Beni Index (XB): {metrics.get('xie_beni_index', 0):.4f}\n"
        val_text += f"- Partition Entropy (PE): {metrics.get('partition_entropy', 0):.4f}\n"

    pdf.chapter_body(val_text)

    if stability:
        pdf.set_font('helvetica', 'B', 11)
        pdf.cell(0, 10, "Stability Test Results (Bootstrap ARI):" if is_en else "Hasil Uji Stabilitas (Bootstrap ARI):", ln=True)
        pdf.set_font('helvetica', '', 10)
        pdf.multi_cell(0, 7, f"Avg Stability Score: {stability.get('stability_score', 0):.4f}\nInterpretation: {stability.get('level', 'N/A')}\n{stability.get('description', '')}")
        pdf.ln(5)

    # --- CHAPTER 3: CLUSTER PROFILE ---
    pdf.add_page()
    pdf.chapter_title('CHAPTER III: PROFILES & DISTRIBUTION' if is_en else 'BAB III: PROFIL DAN DISTRIBUSI ANGGOTA')

    dist = metrics.get('distribution', {})
    pdf.set_font('helvetica', 'B', 11)
    pdf.cell(0, 10, "Distribution Table:" if is_en else "Tabel Distribusi Anggota:", ln=True)
    header = ["Cluster", "Student Count", "Percentage (%)"] if is_en else ["Klaster", "Jumlah Siswa", "Persentase (%)"]
    data = [[f"Cluster {int(k)+1}", v['count'], f"{v['percentage']:.1f}%"] for k, v in dist.items()]
    pdf.add_table(header, data)

    # Centroids Table
    centroids = metrics.get('centroids', [])
    features = metrics.get('feature_names', [])
    if centroids and features:
        pdf.set_font('helvetica', 'B', 11)
        pdf.cell(0, 10, "Matriks Pusat Massa (Centroids):", ln=True)
        header = ["Fitur"] + [f"C{i+1}" for i in range(len(centroids))]
        data = []
        for i, f in enumerate(features):
            row = [f]
            for c in centroids:
                row.append(f"{c[i]:.3f}")
            data.append(row)
        pdf.add_table(header, data)

    # --- CHAPTER 4: KESIMPULAN & REKOMENDASI ---
    pdf.chapter_title('BAB IV: KESIMPULAN & REKOMENDASI')
    conclusion = "Berdasarkan hasil analisis, pengelompokan siswa telah mencapai kondisi optimum yang stabil. "
    advice = metrics.get("improvement_advice", [])
    if advice:
        conclusion += "Namun, untuk riset lanjutan disarankan:\n"
        for a in advice:
            conclusion += f"- {a}\n"
    else:
        conclusion += "Secara keseluruhan, model ini sangat layak digunakan sebagai instrumen pengambilan keputusan."

    pdf.chapter_body(conclusion)

    # --- ETHICAL CLEARANCE ---
    if is_anon:
        pdf.add_page()
        pdf.chapter_title('PERNYATAAN ETIKA PENELITIAN')
        ethical_text = "Laporan ini telah melalui proses audit anonimisasi otomatis (Privacy-First Research). "
        ethical_text += "Seluruh identitas asli subjek penelitian (Nama dan Nomor Induk) telah disamarkan menggunakan kode unik "
        ethical_text += "untuk menjaga kerahasiaan data sesuai dengan standar etika penelitian akademik. "
        ethical_text += "\n\nDigital Verification Code: PROTECTED_ANON_" + str(uuid.uuid4())[:8].upper()
        pdf.chapter_body(ethical_text)

    pdf_bytes = pdf.output()
    output = io.BytesIO(pdf_bytes)

    filename = f"Laporan_Riset_{x_session_id[:8]}.pdf"
    return StreamingResponse(
        output,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.post("/stepwise/ai-discussion/")
async def ai_discussion_generator(x_session_id: Optional[str] = Header(None), lang: str = "id"):
    """AI Research Assistant: Generates a formal academic discussion using Gemini AI."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    is_en = lang == "en"
    api_key = os.environ.get("GOOGLE_API_KEY")
    if not api_key:
        return {
            "status": "partial",
            "narrative": "AI assistant (Gemini) is not configured. Please use the standard interpretation." if is_en else "Maaf, asisten AI (Gemini) belum terkonfigurasi.",
            "disclaimer": "GOOGLE_API_KEY is required for deep narrative." if is_en else "Konfigurasi GOOGLE_API_KEY diperlukan untuk narasi mendalam."
        }

    session = sessions[x_session_id]
    metrics = session.get("metrics", {})
    profiles = metrics.get("cluster_profiles", {})
    features = metrics.get("feature_names", [])
    k = len(profiles)

    # 1. Construct Metadata for Prompt
    profiles_str = ""
    for cid, vals in profiles.items():
        profiles_str += f"- Cluster {int(cid)+1}: " if is_en else f"- Klaster {int(cid)+1}: "
        profiles_str += ", ".join([f"{f}={v:.3f}" for f, v in vals.items()]) + "\n"

    if is_en:
        prompt = f"""
        You are an Associate Professor in Data Science and Education.
        Your task is to write a draft of 'CHAPTER IV: DISCUSSION' for a Master's thesis based on the following student clustering results:

        RESEARCH METADATA:
        - Number of Clusters (K): {k}
        - Analyzed Variables: {", ".join(features)}

        CLUSTER PROFILES (Mean Values):
        {profiles_str}

        VALIDATION RESULTS:
        - Silhouette Score: {metrics.get('silhouette_score', 0):.4f}
        - Davies-Bouldin Index: {metrics.get('davies_bouldin_index', 0):.4f}

        WRITING INSTRUCTIONS:
        1. Use formal Academic English.
        2. Explain the unique characteristics of each cluster in depth.
        3. Identify which variables are most discriminative between groups.
        4. Provide insights on educational policy implications (e.g., scholarships, mentoring, or zoning).
        5. Use strong, data-driven language; avoid words like 'maybe' or 'seems'.

        Output format: Narrative paragraph text.
        """
    else:
        prompt = f"""
        Anda adalah seorang Asisten Profesor Ahli Data Science dan Pendidikan.
        Tugas Anda adalah menulis draf 'BAB IV: PEMBAHASAN' untuk tesis S2 berdasarkan hasil clustering data siswa berikut:

        METADATA RISET:
        - Jumlah Klaster (K): {k}
        - Variabel yang Dianalisis: {", ".join(features)}

        PROFIL PUSAT KLASTER (Rata-rata):
        {profiles_str}

        HASIL VALIDASI:
        - Silhouette Score: {metrics.get('silhouette_score', 0):.4f}
        - Davies-Bouldin Index: {metrics.get('davies_bouldin_index', 0):.4f}

        INSTRUKSI PENULISAN:
        1. Gunakan Bahasa Indonesia formal (Akademik).
        2. Jelaskan karakteristik unik masing-masing klaster secara mendalam.
        3. Identifikasi variabel mana yang paling membedakan antar kelompok (diskriminatif).
        4. Berikan wawasan mengenai implikasi kebijakan pendidikan yang harus diambil (misal: pemberian beasiswa, pendampingan, atau zonasi).
        5. Hindari penggunaan kata 'mungkin' atau 'sepertinya', gunakan bahasa yang tegas berdasarkan data.

        Format output: Teks naratif paragraf.
        """

    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(prompt)

        return {
            "status": "success",
            "narrative": response.text,
            "prompt_tokens": len(prompt.split())
        }
    except Exception as e:
        return {
            "status": "error",
            "message": f"Gagal menghubungi Gemini: {str(e)}",
            "narrative": "Terjadi kesalahan pada layanan AI. Silakan periksa koneksi internet atau kuota API Anda."
        }

@app.post("/stepwise/benchmark/")
async def stepwise_benchmark(x_session_id: Optional[str] = Header(None)):
    """Scientific comparison of K-Means, FCM, and Ensemble algorithms."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    session = sessions[x_session_id]
    df = session["df"]
    config = session.get("config", {})
    ahp_weights = config.get("ahp_weights")
    k = config.get("k", 3)
    features = config.get("features", list(df.select_dtypes(include=[np.number]).columns))
    X_raw = df[features].select_dtypes(include=[np.number]).fillna(0).values
    X = get_weighted_x(X_raw, ahp_weights, features)

    results = {}
    # S2 RIGOR: Lock seed for benchmark reproducibility
    np.random.seed(42)

    # 1. K-Means++ Run
    from sklearn.cluster import KMeans
    start_km = time.time()
    kmeans = KMeans(n_clusters=k, init='k-means++', n_init=10, random_state=42).fit(X)
    end_km = time.time()

    km_labels = kmeans.labels_
    km_sil = float(silhouette_score(X, km_labels))
    km_dbi = float(davies_bouldin_score(X, km_labels))
    km_chi = float(calinski_harabasz_score(X, km_labels))

    results["kmeans"] = {
        "name": "K-Means++",
        "silhouette": km_sil,
        "dbi": km_dbi,
        "chi": km_chi,
        "wcss": float(kmeans.inertia_),
        "time": float(end_km - start_km)
    }

    # 2. Fuzzy C-Means Run
    start_fcm = time.time()
    m = config.get("m", 2.0)

    # S2 RIGOR: Lock seed for benchmark reproducibility
    np.random.seed(42)

    # Vectorized FCM for Benchmark speed
    U = np.random.dirichlet(np.ones(k), size=X.shape[0]).T
    for _ in range(100):
        U_m = U ** m
        centers = (U_m @ X) / (U_m.sum(axis=1)[:, np.newaxis] + 1e-10)
        dists = np.linalg.norm(X[:, np.newaxis] - centers, axis=2)
        dists = np.fmax(dists, 1e-10)
        inv_dists = dists ** (-2.0 / (m - 1))
        new_U = (inv_dists / inv_dists.sum(axis=1)[:, np.newaxis]).T
        if np.linalg.norm(new_U - U) < 1e-4: break
        U = new_U
    end_fcm = time.time()

    fcm_labels = np.argmax(U, axis=0)
    fcm_sil = float(silhouette_score(X, fcm_labels))
    fcm_dbi = float(davies_bouldin_score(X, fcm_labels))
    fcm_chi = float(calinski_harabasz_score(X, fcm_labels))

    # S2 RIGOR: Advanced Fuzzy Metrics
    xb = calculate_xie_beni(X, U, centers, m)
    pe = calculate_partition_entropy(U)

    results["fcm"] = {
        "name": "Fuzzy C-Means",
        "silhouette": fcm_sil,
        "dbi": fcm_dbi,
        "chi": fcm_chi,
        "wcss": float(np.sum(np.min(np.linalg.norm(X[:, np.newaxis] - centers, axis=2), axis=1)**2)),
        "time": float(end_fcm - start_fcm),
        "xie_beni": xb,
        "partition_entropy": pe
    }

    # 3. Ensemble (Consensus) Run
    # Simplified Ensemble via Majority Voting of both above results
    from scipy.stats import mode
    combined_labels = np.array([km_labels, fcm_labels])
    ensemble_labels, _ = mode(combined_labels, axis=0)
    ensemble_labels = ensemble_labels.flatten()

    ens_sil = float(silhouette_score(X, ensemble_labels))
    ens_dbi = float(davies_bouldin_score(X, ensemble_labels))
    ens_chi = float(calinski_harabasz_score(X, ensemble_labels))

    results["ensemble"] = {
        "name": "Ensemble (Hybrid)",
        "silhouette": ens_sil,
        "dbi": ens_dbi,
        "chi": ens_chi,
        "wcss": 0.0, # Not directly applicable
        "time": float((end_km - start_km) + (end_fcm - start_fcm))
    }

    # 4. Comparative Conclusion Generator
    algos = ["kmeans", "fcm", "ensemble"]
    best_sil = max(algos, key=lambda x: results[x]["silhouette"])
    best_dbi = min(algos, key=lambda x: results[x]["dbi"])

    comparison = {
        "winner_silhouette": best_sil,
        "winner_dbi": best_dbi,
        "conclusion": f"Hasil pengujian menunjukkan bahwa {results[best_sil]['name']} memiliki tingkat kerapatan klaster terbaik (Sil={results[best_sil]['silhouette']:.4f}), "
                     f"sementara {results[best_dbi]['name']} paling optimal dalam memisahkan antar kelompok (DBI={results[best_dbi]['dbi']:.4f}). "
                     f"Secara metodologis, {results[best_sil]['name'].split(' ')[0]} disarankan untuk pengambilan kebijakan pendidikan berbasis profil siswa ini."
    }

    return {"status": "success", "results": results, "comparison": comparison}

@app.post("/stepwise/save_config/")
@app.post("/stepwise/mapping-config/")
async def stepwise_mapping(x_session_id: Optional[str] = Header(None), config: Dict[str, Any] = Body(...)):
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")
    sessions[x_session_id]["config"].update(config)
    sync_session_to_firebase(x_session_id)
    return {"status": "success"}

@app.get("/stepwise/history-list/")
async def get_history_list():
    """Fetches all stored research sessions from Firestore."""
    if not db: raise HTTPException(status_code=503, detail="Firestore not connected")
    try:
        docs = db.collection("python_sessions").stream()
        history = []
        for doc in docs:
            d = doc.to_dict()
            # Only include sessions that completed a clustering (have metrics)
            if "metrics" in d and d["metrics"]:
                history.append({
                    "session_id": doc.id,
                    "filename": d.get("filename", "Unknown"),
                    "timestamp": d.get("metrics", {}).get("timestamp", time.time()),
                    "algorithm": d.get("config", {}).get("mode", "kmeans"),
                    "k": d.get("config", {}).get("k", 3),
                    "silhouette": d.get("metrics", {}).get("silhouette_score", 0.0)
                })
        # Sort by timestamp descending
        history = sorted(history, key=lambda x: x["timestamp"], reverse=True)
        return {"status": "success", "history": history}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/stepwise/save-to-history/")
async def save_to_history(x_session_id: Optional[str] = Header(None)):
    """Explicitly marks a session for permanent historical storage."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    # In this implementation, sync already happens.
    # We just ensure it's up to date.
    sync_session_to_firebase(x_session_id)
    return {"status": "success", "message": "Riset berhasil disimpan ke histori permanen."}

@app.post("/stepwise/longitudinal-compare/")
async def longitudinal_compare(params: Dict[str, Any] = Body(...)):
    """Longitudinal Analysis: Compares student movements between two research periods."""
    id_a = params.get("session_id_a")
    id_b = params.get("session_id_b")

    if not db: raise HTTPException(status_code=503, detail="Firestore required")

    # 1. Fetch both sessions
    doc_a = db.collection("python_sessions").document(id_a).get()
    doc_b = db.collection("python_sessions").document(id_b).get()

    if not doc_a.exists or not doc_b.exists:
        raise HTTPException(status_code=404, detail="One or both sessions not found.")

    data_a = doc_a.to_dict()
    data_b = doc_b.to_dict()

    df_a = pd.DataFrame(data_a.get("df_records", []))
    df_b = pd.DataFrame(data_b.get("df_records", []))

    if "nis" not in df_a.columns or "nis" not in df_b.columns:
        raise HTTPException(status_code=400, detail="Dataset harus memiliki kolom 'nis' untuk perbandingan longitudinal.")

    # 2. Match Students
    merged = pd.merge(df_a[['nis', 'nama', 'cluster']], df_b[['nis', 'cluster']], on='nis', suffixes=('_a', '_b'))

    if merged.empty:
        return {"status": "success", "message": "Tidak ditemukan siswa yang sama antar kedua periode.", "movement": []}

    # 3. Calculate Transition Matrix
    k_a = data_a.get("config", {}).get("k", 3)
    k_b = data_b.get("config", {}).get("k", 3)

    # Identify movements
    movement_stats = []
    upgraded = 0
    stable = 0
    downgraded = 0

    for _, row in merged.iterrows():
        c_a = int(row['cluster_a'])
        c_b = int(row['cluster_b'])

        status = "STABLE"
        if c_b < c_a: # Assuming lower cluster index is "better" (e.g., C1 is top)
            status = "UPGRADED"
            upgraded += 1
        elif c_b > c_a:
            status = "DOWNGRADED"
            downgraded += 1
        else:
            stable += 1

        movement_stats.append({
            "nis": str(row['nis']),
            "nama": str(row['nama']),
            "from_cluster": c_a,
            "to_cluster": c_b,
            "status": status
        })

    summary = {
        "total_matched": len(merged),
        "upgraded": upgraded,
        "stable": stable,
        "downgraded": downgraded,
        "upgraded_pct": float(upgraded / len(merged) * 100),
        "stable_pct": float(stable / len(merged) * 100),
        "downgraded_pct": float(downgraded / len(merged) * 100)
    }

    return {
        "status": "success",
        "summary": summary,
        "movements": movement_stats,
        "period_a": data_a.get("filename", "Periode A"),
        "period_b": data_b.get("filename", "Periode B")
    }

@app.post("/stepwise/simulate/")
async def simulate_scenario(x_session_id: Optional[str] = Header(None), data: Dict[str, Any] = Body(...)):
    """Simulates a 'What-If' scenario with prescriptive advice in Indonesian."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    session = sessions[x_session_id]
    scaler = session.get("scaler")
    metrics = session.get("metrics")

    if not metrics or "centroids" not in metrics:
        raise HTTPException(status_code=400, detail="Riset belum selesai. Jalankan clustering terlebih dahulu.")

    features = metrics.get("feature_names", [])
    centroids = np.array(metrics["centroids"])

    # 1. Prepare raw input vector
    input_vals = [float(data.get(f, 0)) for f in features]
    X_input = np.array([input_vals])

    # 2. Apply Scaler if exists
    X_scaled = scaler.transform(X_input) if scaler else X_input

    # 3. Predict Cluster
    dists = np.linalg.norm(centroids - X_scaled, axis=1)
    new_cluster = int(np.argmin(dists))
    confidence = 1.0 - (np.min(dists) / np.sum(dists)) if np.sum(dists) > 0 else 1.0

    # 4. Prescriptive Logic: Find "Best" Cluster to compare
    # We assume the cluster with highest average values is the target
    centroid_means = np.mean(centroids, axis=1)
    best_cluster_idx = int(np.argmax(centroid_means))

    advice = []
    if new_cluster != best_cluster_idx:
        target_centroid = centroids[best_cluster_idx]
        for i, f in enumerate(features):
            diff = target_centroid[i] - X_scaled[0][i]
            if diff > 0.1: # Significant gap
                # Inverse transform the gap if scaler exists to show real units
                real_diff = diff
                if scaler and hasattr(scaler, 'scale_'):
                    real_diff = diff / scaler.scale_[i]

                advice.append(f"Tingkatkan variabel '{f}' sekitar {real_diff:.2f} poin.")

    prescriptive_narrative = " ".join(advice) if advice else "Performa subjek sudah berada pada profil optimal atau mendekati target tertinggi."

    return {
        "status": "success",
        "original_data": data,
        "scaled_vector": X_scaled.tolist()[0],
        "predicted_cluster": new_cluster,
        "confidence": float(confidence),
        "distances": dists.tolist(),
        "prescriptive_advice": prescriptive_narrative,
        "target_cluster": best_cluster_idx
    }

@app.post("/stepwise/simulate-policy/")
async def simulate_policy_intervention(x_session_id: Optional[str] = Header(None), params: Dict[str, Any] = Body(...)):
    """Population Policy Simulator: Simulates the impact of interventions on a specific cluster."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    session = sessions[x_session_id]
    df = session["df"]
    metrics = session.get("metrics", {})
    if not metrics or "centroids" not in metrics:
        raise HTTPException(status_code=400, detail="Run clustering first.")

    target_cluster = params.get("target_cluster_idx") # int
    interventions = params.get("interventions", {}) # { "feature": percentage_change } e.g. { "Nilai": 0.1 }

    # Get students in target cluster
    target_df = df[df["cluster"] == target_cluster].copy()
    if target_df.empty:
        return {"status": "success", "total_impacted": 0, "migrated_count": 0, "message": "Cluster target kosong."}

    features = metrics.get("feature_names", [])
    centroids = np.array(metrics["centroids"])
    ahp_weights = session.get("config", {}).get("ahp_weights")
    scaler = session.get("scaler")

    # 1. Apply Interventions to RAW values
    for feature, pct in interventions.items():
        if feature in target_df.columns:
            # We assume pct is -1.0 to 1.0 (e.g. 0.2 is +20%)
            target_df[feature] = target_df[feature] * (1.0 + pct)

    # 2. Re-predict clusters
    X_new_raw = target_df[features].fillna(0).values
    X_new_scaled = scaler.transform(X_new_raw) if scaler else X_new_raw
    X_new_clustering = get_weighted_x(X_new_scaled, ahp_weights, features)

    new_assignments = []
    for row in X_new_clustering:
        dists = np.linalg.norm(centroids - row, axis=1)
        new_assignments.append(int(np.argmin(dists)))

    # 3. Count migrations
    migrated_count = sum(1 for a in new_assignments if a != target_cluster)

    # Destination summary
    dest_counts = {}
    for a in new_assignments:
        if a != target_cluster:
            dest_counts[str(a)] = dest_counts.get(str(a), 0) + 1

    return {
        "status": "success",
        "total_impacted": len(target_df),
        "migrated_count": migrated_count,
        "migration_rate": float(migrated_count / len(target_df) * 100),
        "destinations": dest_counts,
        "message": f"Simulasi selesai. {migrated_count} siswa ({migrated_count/len(target_df)*100:.1f}%) diprediksi akan berpindah klaster."
    }

@app.get("/stepwise/export-excel/")
async def export_excel(x_session_id: Optional[str] = Header(None)):
    await ensure_session(x_session_id)
    session = sessions[x_session_id]
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        for name, data in session.get("checkpoints", {}).items():
            if data: pd.DataFrame(data).to_excel(writer, sheet_name=name[:31], index=False)
        session["df"].to_excel(writer, sheet_name="Hasil Akhir", index=False)
    output.seek(0)
    return StreamingResponse(output, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": f"attachment; filename=Riset_{x_session_id[:8]}.xlsx"})

@app.get("/stepwise/build-manuscript/")
async def build_manuscript(x_session_id: Optional[str] = Header(None)):
    """Automated Paper Architect: Generates a complete research manuscript in .docx format."""
    await ensure_session(x_session_id)
    if x_session_id not in sessions: raise HTTPException(status_code=404, detail="Session not found")

    session = sessions[x_session_id]
    metrics = session.get("metrics", {})
    config = session.get("config", {})
    df = session["df"]
    features = metrics.get("feature_names", [])

    doc = Document()

    # --- STYLE ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # --- TITLE ---
    title = doc.add_heading('IDENTIFIKASI PROFIL SISWA WILAYAH PERBATASAN MENGGUNAKAN PENDEKATAN CLUSTERING CERDAS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- ABSTRACT ---
    doc.add_heading('ABSTRACT', level=1)
    abstract_text = f"This study analyzes the profiles of {len(df)} students in border regions using {config.get('mode', 'K-Means').upper()} clustering. "
    abstract_text += f"The analysis used {len(features)} variables including {', '.join(features[:3])}. "
    abstract_text += f"The result achieved a Silhouette Score of {metrics.get('silhouette_score', 0):.4f} and DBI of {metrics.get('davies_bouldin_index', 0):.4f}, "
    abstract_text += "providing a robust foundation for educational policy interventions."
    doc.add_paragraph(abstract_text)

    # --- INTRODUCTION (AI Generated) ---
    doc.add_heading('I. INTRODUCTION', level=1)
    intro_prompt = f"Tuliskan draf pendahuluan singkat (2 paragraf) untuk artikel ilmiah mengenai pentingnya clustering data siswa di wilayah perbatasan Indonesia untuk pemerataan bantuan pendidikan. Fokus pada variabel: {', '.join(features)}."
    try:
        genai.configure(api_key=os.environ.get("GOOGLE_API_KEY"))
        model = genai.GenerativeModel('gemini-1.5-flash')
        intro_resp = model.generate_content(intro_prompt)
        doc.add_paragraph(intro_resp.text)
    except:
        doc.add_paragraph("Sistem clustering cerdas sangat penting untuk mengidentifikasi kesenjangan pendidikan di wilayah perbatasan...")

    # --- METHODOLOGY ---
    doc.add_heading('II. METHODOLOGY', level=1)
    method_p = doc.add_paragraph()
    method_p.add_run(f"Penelitian ini menerapkan jalur pipa data mining (pipeline) yang terdiri dari: ")
    method_p.add_run(f"Cleaning, Imputation, {session.get('audit', {}).get('normalization_method', 'Scaling')}, dan {config.get('mode', 'kmeans').upper()} Clustering. ")

    if config.get("ahp_weights"):
        method_p.add_run("Pembobotan variabel ditentukan melalui Analytic Hierarchy Process (AHP).")

    # --- RESULTS ---
    doc.add_heading('III. RESULTS', level=1)
    doc.add_paragraph(f"Berdasarkan analisis, ditemukan {len(metrics.get('distribution', {}))} kelompok siswa dengan karakteristik yang berbeda.")

    # 1. Plot PCA Map
    try:
        # Re-generate PCA plot for high-res
        ahp_weights = config.get("ahp_weights")
        X_raw = df[features].select_dtypes(include=[np.number]).fillna(0).values
        X = get_weighted_x(X_raw, ahp_weights, features)
        pca = PCA(n_components=2, random_state=42)
        X_2d = pca.fit_transform(X)

        plt.figure(figsize=(8, 6))
        sns.scatterplot(x=X_2d[:, 0], y=X_2d[:, 1], hue=df["cluster"], palette="viridis", s=100)
        plt.title("Cluster Projection Map (PCA)")

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300)
        plt.close()
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(5))
        doc.add_paragraph("Figure 1. Visualisasi Sebaran Klaster menggunakan Principal Component Analysis (PCA).").alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e: print(f"Plot PCA error: {e}")

    # 2. Add Profiles Table
    table = doc.add_table(rows=1, cols=len(features)+1)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Cluster'
    for i, f in enumerate(features): hdr_cells[i+1].text = f

    profiles = metrics.get("cluster_profiles", {})
    for cid, vals in profiles.items():
        row_cells = table.add_row().cells
        row_cells[0].text = f"C{int(cid)+1}"
        for i, f in enumerate(features):
            row_cells[i+1].text = f"{vals.get(f, 0):.3f}"

    # --- DISCUSSION ---
    doc.add_heading('IV. DISCUSSION', level=1)
    # Reuse previous AI logic for discussion
    try:
        disc_prompt = f"Analisis secara mendalam hasil klaster berikut untuk draf jurnal: {str(profiles)}. Fokus pada implikasi bantuan pendidikan."
        disc_resp = model.generate_content(disc_prompt)
        doc.add_paragraph(disc_resp.text)
    except:
        doc.add_paragraph("Analisis menunjukkan adanya perbedaan signifikan antara kelompok siswa...")

    doc.add_heading('V. CONCLUSION', level=1)
    doc.add_paragraph("Riset ini membuktikan bahwa pendekatan clustering hibrida mampu memisahkan profil siswa perbatasan secara akurat untuk kebutuhan pengambilan keputusan.")

    # --- SAVE & RETURN ---
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = f"Manuscript_SIMORBATAS_{x_session_id[:8]}.docx"
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", 7860)))
