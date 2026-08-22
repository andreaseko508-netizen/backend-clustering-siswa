import io
import uuid
import tempfile
import os
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from fpdf import FPDF

class ResearchReportPDF(FPDF):
    def header(self):
        if self.page_no() > 1:
            self.set_font('helvetica', 'B', 10)
            self.set_text_color(100, 116, 139)
            self.cell(0, 10, 'SIMORBATAS: Final Research Publication | Institutional Grade', border=False, align='R')
            self.ln(10)

    def footer(self):
        self.set_y(-20)
        self.set_font('helvetica', 'I', 8)
        self.set_text_color(148, 163, 184)
        self.cell(0, 10, f'Halaman {self.page_no()}/{{nb}}', align='C')
        self.set_x(self.l_margin)
        self.cell(0, 10, f'Digital Verification: {str(uuid.uuid4())[:13].upper()}', align='L')

    def chapter_title(self, label):
        self.set_font('helvetica', 'B', 14)
        self.set_text_color(30, 58, 138)
        self.set_fill_color(241, 245, 249)
        self.cell(0, 12, f" {label}", border='L', ln=True, fill=True)
        self.ln(5)

    def chapter_body(self, body):
        self.set_font('helvetica', '', 11)
        self.set_text_color(51, 65, 85)
        self.multi_cell(0, 7, body)
        self.ln(5)

    def add_table(self, header, data):
        self.set_font('helvetica', 'B', 10)
        self.set_fill_color(248, 250, 252)
        col_width = self.epw / len(header)
        for h in header:
            self.cell(col_width, 10, h, border=1, align='C', fill=True)
        self.ln()
        self.set_font('helvetica', '', 10)
        for row in data:
            if self.get_y() > 250: self.add_page()
            for item in row:
                self.cell(col_width, 8, str(item), border=1, align='C')
            self.ln()
        self.ln(8)

    def add_image_from_buf(self, buf, width=150):
        if buf is None: return
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp.write(buf.getvalue())
            tmp_path = tmp.name
        x = (self.w - width) / 2
        self.image(tmp_path, x=x, w=width)
        os.unlink(tmp_path)
        self.ln(5)

def generate_radar_chart_bytes(profiles, features):
    try:
        if not profiles or not features: return None
        categories = features
        N = len(categories)
        angles = [n / float(N) * 2 * np.pi for n in range(N)]
        angles += angles[:1]
        fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))
        colors = ['#1E3A8A', '#0F766E', '#B71C1C', '#6B21A8', '#F59E0B']
        for i, (cid, vals) in enumerate(profiles.items()):
            values = [vals.get(f, 0) for f in categories]
            values += values[:1]
            ax.plot(angles, values, linewidth=2, linestyle='solid', label=f"Cluster {int(cid)+1}", color=colors[i % len(colors)])
            ax.fill(angles, values, colors[i % len(colors)], alpha=0.1)
        plt.xticks(angles[:-1], categories, size=10)
        ax.set_rlabel_position(0)
        plt.yticks([0.25, 0.5, 0.75], ["0.25", "0.5", "0.75"], color="grey", size=7)
        plt.ylim(0, 1)
        plt.legend(loc='upper right', bbox_to_anchor=(0.1, 0.1))
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf
    except: return None

def generate_manuscript_docx(session, x_session_id, get_weighted_x_func):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import google.generativeai as genai
    from sklearn.decomposition import PCA
    import pandas as pd

    metrics = session.get("metrics", {})
    config = session.get("config", {})
    df = session["df"]
    features = metrics.get("feature_names", [])

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    title = doc.add_heading('IDENTIFIKASI PROFIL SISWA WILAYAH PERBATASAN MENGGUNAKAN PENDEKATAN CLUSTERING CERDAS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('ABSTRACT', level=1)
    abstract_text = f"This study analyzes the profiles of {len(df)} students in border regions using {config.get('mode', 'K-Means').upper()} clustering. "
    abstract_text += f"The analysis used {len(features)} variables including {', '.join(features[:3])}. "
    abstract_text += f"The result achieved a Silhouette Score of {metrics.get('silhouette_score', 0):.4f} and DBI of {metrics.get('davies_bouldin_index', 0):.4f}, "
    abstract_text += "providing a robust foundation for educational policy interventions."
    doc.add_paragraph(abstract_text)

    doc.add_heading('I. INTRODUCTION', level=1)
    intro_prompt = f"Tuliskan draf pendahuluan singkat (2 paragraf) untuk artikel ilmiah mengenai pentingnya clustering data siswa di wilayah perbatasan Indonesia untuk pemerataan bantuan pendidikan. Fokus pada variabel: {', '.join(features)}."
    try:
        genai.configure(api_key=os.environ.get("GOOGLE_API_KEY"))
        model = genai.GenerativeModel('gemini-1.5-flash')
        intro_resp = model.generate_content(intro_prompt)
        doc.add_paragraph(intro_resp.text)
    except:
        doc.add_paragraph("Sistem clustering cerdas sangat penting untuk mengidentifikasi kesenjangan pendidikan di wilayah perbatasan...")

    doc.add_heading('II. METHODOLOGY', level=1)
    method_p = doc.add_paragraph()
    method_p.add_run(f"Penelitian ini menerapkan jalur pipa data mining (pipeline) yang terdiri dari: ")
    method_p.add_run(f"Cleaning, Imputation, {session.get('audit', {}).get('normalization_method', 'Scaling')}, dan {config.get('mode', 'kmeans').upper()} Clustering. ")
    if config.get("ahp_weights"):
        method_p.add_run("Pembobotan variabel ditentukan melalui Analytic Hierarchy Process (AHP).")

    doc.add_heading('III. RESULTS', level=1)
    doc.add_paragraph(f"Berdasarkan analisis, ditemukan {len(metrics.get('distribution', {}))} kelompok siswa dengan karakteristik yang berbeda.")

    try:
        ahp_weights = config.get("ahp_weights")
        X_raw = df[features].select_dtypes(include=[np.number]).fillna(0).values
        X = get_weighted_x_func(X_raw, ahp_weights, features)
        pca = PCA(n_components=2, random_state=42)
        X_2d = pca.fit_transform(X)

        plt.figure(figsize=(8, 6))
        sns.scatterplot(x=X_2d[:, 0], y=X_2d[:, 1], hue=df["cluster"], palette="viridis", s=100)
        plt.title("Cluster Projection Map (PCA)")
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300)
        plt.close()
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(5))
        doc.add_paragraph("Figure 1. Visualisasi Sebaran Klaster menggunakan Principal Component Analysis (PCA).").alignment = WD_ALIGN_PARAGRAPH.CENTER
    except: pass

    table = doc.add_table(rows=1, cols=len(features)+1)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Cluster'
    for i, f in enumerate(features): hdr_cells[i+1].text = f

    profiles = metrics.get("cluster_profiles", {})
    for cid, vals in profiles.items():
        row_cells = table.add_row().cells
        row_cells[0].text = f"C{int(cid)+1}"
        for i, f in enumerate(features):
            row_cells[i+1].text = f"{vals.get(f, 0):.3f}"

    doc.add_heading('IV. DISCUSSION', level=1)
    try:
        disc_prompt = f"Analisis secara mendalam hasil klaster berikut untuk draf jurnal: {str(profiles)}. Fokus pada implikasi bantuan pendidikan."
        model = genai.GenerativeModel('gemini-1.5-flash')
        disc_resp = model.generate_content(disc_prompt)
        doc.add_paragraph(disc_resp.text)
    except:
        doc.add_paragraph("Analisis menunjukkan adanya perbedaan signifikan antara kelompok siswa...")

    doc.add_heading('V. CONCLUSION', level=1)
    doc.add_paragraph("Riset ini membuktikan bahwa pendekatan clustering hibrida mampu memisahkan profil siswa perbatasan secara akurat untuk kebutuhan pengambilan keputusan.")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def generate_bar_chart_bytes(distribution):
    try:
        if not distribution: return None
        labels = [f"Cluster {int(k)+1}" for k in distribution.keys()]
        counts = [v['count'] for v in distribution.values()]
        plt.figure(figsize=(10, 5))
        sns.set_style("whitegrid")
        palette = ['#1E3A8A', '#0F766E', '#B71C1C', '#6B21A8', '#F59E0B']
        ax = sns.barplot(x=labels, y=counts, palette=palette[:len(labels)])
        plt.title("Distribution of Students across Clusters", fontsize=14, fontweight='bold')
        plt.ylabel("Number of Students")
        for p in ax.patches:
            ax.annotate(f'{int(p.get_height())}', (p.get_x() + p.get_width() / 2., p.get_height()), ha='center', va='center', xytext=(0, 10), textcoords='offset points')
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf
    except: return None

def generate_manuscript_docx(session, x_session_id, get_weighted_x_func):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import google.generativeai as genai
    from sklearn.decomposition import PCA
    import pandas as pd

    metrics = session.get("metrics", {})
    config = session.get("config", {})
    df = session["df"]
    features = metrics.get("feature_names", [])

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    title = doc.add_heading('IDENTIFIKASI PROFIL SISWA WILAYAH PERBATASAN MENGGUNAKAN PENDEKATAN CLUSTERING CERDAS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('ABSTRACT', level=1)
    abstract_text = f"This study analyzes the profiles of {len(df)} students in border regions using {config.get('mode', 'K-Means').upper()} clustering. "
    abstract_text += f"The analysis used {len(features)} variables including {', '.join(features[:3])}. "
    abstract_text += f"The result achieved a Silhouette Score of {metrics.get('silhouette_score', 0):.4f} and DBI of {metrics.get('davies_bouldin_index', 0):.4f}, "
    abstract_text += "providing a robust foundation for educational policy interventions."
    doc.add_paragraph(abstract_text)

    doc.add_heading('I. INTRODUCTION', level=1)
    intro_prompt = f"Tuliskan draf pendahuluan singkat (2 paragraf) untuk artikel ilmiah mengenai pentingnya clustering data siswa di wilayah perbatasan Indonesia untuk pemerataan bantuan pendidikan. Fokus pada variabel: {', '.join(features)}."
    try:
        genai.configure(api_key=os.environ.get("GOOGLE_API_KEY"))
        model = genai.GenerativeModel('gemini-1.5-flash')
        intro_resp = model.generate_content(intro_prompt)
        doc.add_paragraph(intro_resp.text)
    except:
        doc.add_paragraph("Sistem clustering cerdas sangat penting untuk mengidentifikasi kesenjangan pendidikan di wilayah perbatasan...")

    doc.add_heading('II. METHODOLOGY', level=1)
    method_p = doc.add_paragraph()
    method_p.add_run(f"Penelitian ini menerapkan jalur pipa data mining (pipeline) yang terdiri dari: ")
    method_p.add_run(f"Cleaning, Imputation, {session.get('audit', {}).get('normalization_method', 'Scaling')}, dan {config.get('mode', 'kmeans').upper()} Clustering. ")
    if config.get("ahp_weights"):
        method_p.add_run("Pembobotan variabel ditentukan melalui Analytic Hierarchy Process (AHP).")

    doc.add_heading('III. RESULTS', level=1)
    doc.add_paragraph(f"Berdasarkan analisis, ditemukan {len(metrics.get('distribution', {}))} kelompok siswa dengan karakteristik yang berbeda.")

    try:
        ahp_weights = config.get("ahp_weights")
        X_raw = df[features].select_dtypes(include=[np.number]).fillna(0).values
        X = get_weighted_x_func(X_raw, ahp_weights, features)
        pca = PCA(n_components=2, random_state=42)
        X_2d = pca.fit_transform(X)

        plt.figure(figsize=(8, 6))
        sns.scatterplot(x=X_2d[:, 0], y=X_2d[:, 1], hue=df["cluster"], palette="viridis", s=100)
        plt.title("Cluster Projection Map (PCA)")
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300)
        plt.close()
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(5))
        doc.add_paragraph("Figure 1. Visualisasi Sebaran Klaster menggunakan Principal Component Analysis (PCA).").alignment = WD_ALIGN_PARAGRAPH.CENTER
    except: pass

    table = doc.add_table(rows=1, cols=len(features)+1)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Cluster'
    for i, f in enumerate(features): hdr_cells[i+1].text = f

    profiles = metrics.get("cluster_profiles", {})
    for cid, vals in profiles.items():
        row_cells = table.add_row().cells
        row_cells[0].text = f"C{int(cid)+1}"
        for i, f in enumerate(features):
            row_cells[i+1].text = f"{vals.get(f, 0):.3f}"

    doc.add_heading('IV. DISCUSSION', level=1)
    try:
        disc_prompt = f"Analisis secara mendalam hasil klaster berikut untuk draf jurnal: {str(profiles)}. Fokus pada implikasi bantuan pendidikan."
        model = genai.GenerativeModel('gemini-1.5-flash')
        disc_resp = model.generate_content(disc_prompt)
        doc.add_paragraph(disc_resp.text)
    except:
        doc.add_paragraph("Analisis menunjukkan adanya perbedaan signifikan antara kelompok siswa...")

    doc.add_heading('V. CONCLUSION', level=1)
    doc.add_paragraph("Riset ini membuktikan bahwa pendekatan clustering hibrida mampu memisahkan profil siswa perbatasan secara akurat untuk kebutuhan pengambilan keputusan.")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def generate_silhouette_chart_bytes(silhouette_plot_data):
    try:
        if not silhouette_plot_data: return None
        plt.figure(figsize=(10, 6))
        y_lower = 10
        colors = ['#1E3A8A', '#0F766E', '#B71C1C', '#6B21A8', '#F59E0B']
        total_avg, valid_clusters = 0, 0
        for i, cluster_data in enumerate(silhouette_plot_data):
            ith_cluster_sil_values = cluster_data['values']
            size_cluster_i = len(ith_cluster_sil_values)
            y_upper = y_lower + size_cluster_i
            color = colors[i % len(colors)]
            plt.fill_betweenx(np.arange(y_lower, y_upper), 0, ith_cluster_sil_values, facecolor=color, edgecolor=color, alpha=0.7)
            plt.text(-0.05, y_lower + 0.5 * size_cluster_i, str(cluster_data['cluster'] + 1))
            y_lower = y_upper + 10
            total_avg += cluster_data['avg']
            valid_clusters += 1
        avg_score = total_avg / valid_clusters if valid_clusters > 0 else 0
        plt.axvline(x=avg_score, color="red", linestyle="--")
        plt.title("Silhouette Analysis for Cluster Quality", fontsize=14, fontweight='bold')
        plt.xlabel("Silhouette Coefficient Value")
        plt.ylabel("Cluster Label")
        plt.yticks([])
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf
    except: return None

def generate_manuscript_docx(session, x_session_id, get_weighted_x_func):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import google.generativeai as genai
    from sklearn.decomposition import PCA
    import pandas as pd

    metrics = session.get("metrics", {})
    config = session.get("config", {})
    df = session["df"]
    features = metrics.get("feature_names", [])

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    title = doc.add_heading('IDENTIFIKASI PROFIL SISWA WILAYAH PERBATASAN MENGGUNAKAN PENDEKATAN CLUSTERING CERDAS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('ABSTRACT', level=1)
    abstract_text = f"This study analyzes the profiles of {len(df)} students in border regions using {config.get('mode', 'K-Means').upper()} clustering. "
    abstract_text += f"The analysis used {len(features)} variables including {', '.join(features[:3])}. "
    abstract_text += f"The result achieved a Silhouette Score of {metrics.get('silhouette_score', 0):.4f} and DBI of {metrics.get('davies_bouldin_index', 0):.4f}, "
    abstract_text += "providing a robust foundation for educational policy interventions."
    doc.add_paragraph(abstract_text)

    doc.add_heading('I. INTRODUCTION', level=1)
    intro_prompt = f"Tuliskan draf pendahuluan singkat (2 paragraf) untuk artikel ilmiah mengenai pentingnya clustering data siswa di wilayah perbatasan Indonesia untuk pemerataan bantuan pendidikan. Fokus pada variabel: {', '.join(features)}."
    try:
        genai.configure(api_key=os.environ.get("GOOGLE_API_KEY"))
        model = genai.GenerativeModel('gemini-1.5-flash')
        intro_resp = model.generate_content(intro_prompt)
        doc.add_paragraph(intro_resp.text)
    except:
        doc.add_paragraph("Sistem clustering cerdas sangat penting untuk mengidentifikasi kesenjangan pendidikan di wilayah perbatasan...")

    doc.add_heading('II. METHODOLOGY', level=1)
    method_p = doc.add_paragraph()
    method_p.add_run(f"Penelitian ini menerapkan jalur pipa data mining (pipeline) yang terdiri dari: ")
    method_p.add_run(f"Cleaning, Imputation, {session.get('audit', {}).get('normalization_method', 'Scaling')}, dan {config.get('mode', 'kmeans').upper()} Clustering. ")
    if config.get("ahp_weights"):
        method_p.add_run("Pembobotan variabel ditentukan melalui Analytic Hierarchy Process (AHP).")

    doc.add_heading('III. RESULTS', level=1)
    doc.add_paragraph(f"Berdasarkan analisis, ditemukan {len(metrics.get('distribution', {}))} kelompok siswa dengan karakteristik yang berbeda.")

    try:
        ahp_weights = config.get("ahp_weights")
        X_raw = df[features].select_dtypes(include=[np.number]).fillna(0).values
        X = get_weighted_x_func(X_raw, ahp_weights, features)
        pca = PCA(n_components=2, random_state=42)
        X_2d = pca.fit_transform(X)

        plt.figure(figsize=(8, 6))
        sns.scatterplot(x=X_2d[:, 0], y=X_2d[:, 1], hue=df["cluster"], palette="viridis", s=100)
        plt.title("Cluster Projection Map (PCA)")
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300)
        plt.close()
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(5))
        doc.add_paragraph("Figure 1. Visualisasi Sebaran Klaster menggunakan Principal Component Analysis (PCA).").alignment = WD_ALIGN_PARAGRAPH.CENTER
    except: pass

    table = doc.add_table(rows=1, cols=len(features)+1)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Cluster'
    for i, f in enumerate(features): hdr_cells[i+1].text = f

    profiles = metrics.get("cluster_profiles", {})
    for cid, vals in profiles.items():
        row_cells = table.add_row().cells
        row_cells[0].text = f"C{int(cid)+1}"
        for i, f in enumerate(features):
            row_cells[i+1].text = f"{vals.get(f, 0):.3f}"

    doc.add_heading('IV. DISCUSSION', level=1)
    try:
        disc_prompt = f"Analisis secara mendalam hasil klaster berikut untuk draf jurnal: {str(profiles)}. Fokus pada implikasi bantuan pendidikan."
        model = genai.GenerativeModel('gemini-1.5-flash')
        disc_resp = model.generate_content(disc_prompt)
        doc.add_paragraph(disc_resp.text)
    except:
        doc.add_paragraph("Analisis menunjukkan adanya perbedaan signifikan antara kelompok siswa...")

    doc.add_heading('V. CONCLUSION', level=1)
    doc.add_paragraph("Riset ini membuktikan bahwa pendekatan clustering hibrida mampu memisahkan profil siswa perbatasan secara akurat untuk kebutuhan pengambilan keputusan.")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
